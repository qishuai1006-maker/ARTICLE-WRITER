import sys
import subprocess

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

import re
import os

md_file = "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/Super Writer/outputs/midea_ac_20260617/03_最终成片.md"
cover_img = "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/Super Writer/outputs/midea_ac_20260617/封面.png"
info_img1 = "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/Super Writer/outputs/midea_ac_20260617/04_信息图1_外机参数避坑.png"
output_docx = "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/Super Writer/outputs/midea_ac_20260617/美的空调长文_配图排版版.docx"

doc = Document()

# Set styles
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = '黑体'
h1_font.size = Pt(20)

h2_style = doc.styles['Heading 2']
h2_font = h2_style.font
h2_font.name = '黑体'
h2_font.size = Pt(16)

with open(md_file, "r", encoding="utf-8") as f:
    lines = f.readlines()

def add_paragraph_with_bold(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    # Split by **...**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

for i, line in enumerate(lines):
    line = line.strip()
    if not line:
        continue
        
    if line.startswith('# '):
        p = doc.add_paragraph(line[2:], style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add cover image after title
        if os.path.exists(cover_img):
            doc.add_picture(cover_img, width=Inches(6.0))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    elif line.startswith('## '):
        doc.add_paragraph(line[3:], style='Heading 2')
        
    elif line.startswith('[此处插入信息图1'):
        if os.path.exists(info_img1):
            doc.add_picture(info_img1, width=Inches(6.0))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            add_paragraph_with_bold(doc, "[信息图1缺失]")
            
    elif line.startswith('[此处插入信息图2'):
        # Skip or add note
        pass
        
    elif line.startswith('- '):
        add_paragraph_with_bold(doc, line, style='List Bullet')
        
    elif line.startswith('---'):
        doc.add_page_break()
        
    elif line.startswith('|'):
        # Just simple text representation for table or skip if it's the QC table
        if "质检维度" in line or "---" in line or "标题拦截力" in line:
            add_paragraph_with_bold(doc, line)
        else:
            add_paragraph_with_bold(doc, line)
    else:
        add_paragraph_with_bold(doc, line)

doc.save(output_docx)
print(f"Docx saved to {output_docx}")
