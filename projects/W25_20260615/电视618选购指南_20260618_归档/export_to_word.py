#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""export_to_word.py — ZCode 本地修正版图文 Word 导出。
修正 Scripts/build_article_docx.py 的 set_font 参数错位 bug（第96/105行）。
仅用于 outputs/zcode/，不改原脚本。原脚本修好后此文件可删。

用法:
  python3 export_to_word.py \
    --md outputs/zcode/03_风控主编终稿_冰箱.md \
    --out outputs/zcode/海尔冰箱单系统猫腻_图文稿_20260618.docx \
    --img1 outputs/zcode/T5_双系统对比.png \
    --img2 outputs/zcode/T5_海尔型号清单.png
"""
import argparse, re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CN_BODY, CN_HEAD = "宋体", "黑体"
IMG_W = 15.5


def set_font(run, size=11, bold=False, color=None, name_cn=CN_BODY):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_cn
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_runs_bold(para, text, size=11, name_cn=CN_BODY, base_bold=False):
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            set_font(r, size, True, None, name_cn)
        else:
            r = para.add_run(part)
            set_font(r, size, base_bold, None, name_cn)


def add_pic(doc, path, w=IMG_W):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Cm(w))


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, 10, True, (0xC0, 0x00, 0x00))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--md', required=True)
    ap.add_argument('--out', required=True)
    ap.add_argument('--cover')
    ap.add_argument('--img1')
    ap.add_argument('--img2')
    ap.add_argument('--img1-kw', default='信息图1')
    ap.add_argument('--img2-kw', default='信息图2')
    a = ap.parse_args()

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.5)

    def maybe_pic(path):
        if path and os.path.exists(path):
            add_pic(doc, path)
        else:
            add_placeholder(doc, f'【此处插图缺失：{path or "未指定"}】')

    with open(a.md, encoding='utf-8') as f:
        lines = f.read().splitlines()

    need_cover = True
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith('# ') and not s.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            add_runs_bold(p, s[2:].strip(), 18, CN_HEAD, True)
            continue
        if need_cover:
            maybe_pic(a.cover)
            need_cover = False
        if s.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            add_runs_bold(p, s[3:].strip(), 13, CN_HEAD, True)
        elif a.img1_kw in s or '信息图1' in s:
            maybe_pic(a.img1)
        elif a.img2_kw in s or '信息图2' in s:
            maybe_pic(a.img2)
        elif s.startswith('- '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run('· ')
            set_font(r, 11)
            add_runs_bold(p, s[2:].strip(), 11)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(0.74)
            add_runs_bold(p, s, 11)

    doc.save(a.out)
    print('saved:', a.out)


if __name__ == '__main__':
    main()
