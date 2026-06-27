import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # 修改默认中文字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(12)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # 遇到质检表就停止，不输出内部打分到 Word 里
        if '附：系统强制执行的 100 分爆款质检表' in line:
            break
            
        # 匹配图片语法 ![alt](path)
        img_match = re.search(r'!\[.*?\]\((.*?)\)', line)
        if img_match:
            img_path = img_match.group(1)
            # 如果是封面图，强制替换为刚才生成的最新版实拍图
            if '封面图' in img_path:
                img_path = '封面图_最新版.png'
                
            img_path = os.path.join(os.path.dirname(md_path), img_path.replace('./', ''))
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # 插入图片并控制最大宽度
                run.add_picture(img_path, width=Inches(6.0))
            continue
            
        # 标题处理
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.replace('# ', '').replace('**', ''))
            run.bold = True
            run.font.size = Pt(18)
            continue
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('### ', '').replace('**', ''))
            run.bold = True
            run.font.size = Pt(14)
            continue
            
        # 正文内容清理星号等Markdown标记
        clean_text = line.replace('**', '').replace('---', '')
        if clean_text:
            p = doc.add_paragraph()
            run = p.add_run(clean_text)
            
    doc.save(docx_path)
    print(f"SUCCESS: Saved to {docx_path}")

if __name__ == "__main__":
    md_file = '/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/Super Writer/outputs/终稿_冰箱_头条.md'
    docx_file = '/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/Super Writer/outputs/终稿_冰箱闭眼入_可直接发布版.docx'
    markdown_to_docx(md_file, docx_file)
