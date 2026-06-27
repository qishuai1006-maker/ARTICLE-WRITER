from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs" / "Codex"
ARTICLE = OUT / "03_风控主编终稿_油烟机.md"
COVER = OUT / "T5_ChatGPT_点击_油烟机封面.png"
INFO_1 = OUT / "T5_收藏_油烟机4参数.png"
INFO_2 = OUT / "T5_对比_油烟机价位段.png"
DOCX = OUT / "牛科技说_油烟机图文稿_含配图_20260618.docx"


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_image(doc, image_path, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    set_paragraph_spacing(paragraph, before=6, after=3)

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, before=0, after=12)
    for run in cap.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=8)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)


def main():
    raw = ARTICLE.read_text(encoding="utf-8").splitlines()
    title = raw[0].lstrip("# ").strip()
    body_lines = raw[2:]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.15

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title_p, before=0, after=3)
    title_run = title_p.add_run(title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    meta = doc.add_paragraph("目标账号：牛科技说｜品类：油烟机｜用途：今日头条图文发布稿")
    set_paragraph_spacing(meta, before=0, after=12)
    for run in meta.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(85, 85, 85)

    add_image(doc, COVER, "封面图：油烟机与中式爆炒场景")

    for line in body_lines:
        text = line.strip()
        if not text:
            continue
        if text == "[此处插入信息图1：油烟机4个参数决策卡]":
            add_image(doc, INFO_1, "信息图：油烟机4个参数决策卡")
            continue
        if text == "[此处插入信息图2：5个价位段代表型号清单]":
            add_image(doc, INFO_2, "信息图：5个价位段代表型号清单")
            continue
        add_body_paragraph(doc, text)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
