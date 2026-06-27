#!/usr/bin/env python3
"""Export the 空调三要三不要 final markdown to a formatted Word doc with all 3 images embedded.

Self-contained: paths resolve relative to this script's directory, so it can
live inside the archive folder. Mirrors the create_word.py flow but fixes the
info-image-2 handling (the root script skipped it) and points at this article.
"""

import os
import re
import subprocess
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

DIR = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(DIR, "03_风控主编终稿_空调.md")
COVER = os.path.join(DIR, "T5_ChatGPT_点击.png")
INFO1 = os.path.join(DIR, "T5_收藏.png")
INFO2 = os.path.join(DIR, "T5_对比.png")
OUT = os.path.join(DIR, "空调三要三不要_图文稿_20260618.docx")

doc = Document()
doc.styles["Normal"].font.name = "宋体"
doc.styles["Normal"].font.size = Pt(12)
doc.styles["Heading 1"].font.name = "黑体"
doc.styles["Heading 1"].font.size = Pt(20)
doc.styles["Heading 2"].font.name = "黑体"
doc.styles["Heading 2"].font.size = Pt(16)


def add_paragraph(text, style="Normal"):
    p = doc.add_paragraph(style=style)
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def add_image(path):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_paragraph(f"[图片缺失：{os.path.basename(path)}]")


with open(MD, encoding="utf-8") as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    if line.startswith("# "):
        p = doc.add_paragraph(line[2:], style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_image(COVER)  # cover right under the title
    elif line.startswith("## "):
        doc.add_paragraph(line[3:], style="Heading 2")
    elif line.startswith("[此处插入信息图1"):
        add_image(INFO1)
    elif line.startswith("[此处插入信息图2"):
        add_image(INFO2)
    elif line.startswith("- "):
        add_paragraph(line, style="List Bullet")
    elif line.startswith("---"):
        doc.add_page_break()
    else:
        add_paragraph(line)

doc.save(OUT)
print(f"Docx saved: {OUT}")
