#!/usr/bin/env python3
"""Generate a Word document from T3_头条.md with embedded images."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re
import os

BASE = os.path.dirname(os.path.abspath(__file__))

# Read the markdown article (lines 1-163, before the T3 checklist)
with open(os.path.join(BASE, 'T3_头条.md'), 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Find where the T3 quality checklist starts
body_lines = []
for line in lines:
    if line.strip().startswith('## T3 质量门禁') or line.strip().startswith('---') and 'T3 头条图文' in line:
        break
    body_lines.append(line)

# Remove trailing --- and empty lines at end
while body_lines and (body_lines[-1].strip() == '' or body_lines[-1].strip() == '---'):
    body_lines.pop()

article = ''.join(body_lines)

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Set narrow margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_paragraph_with_bold(doc, text, style_name='Normal'):
    """Add a paragraph, handling **bold** markers."""
    p = doc.add_paragraph()
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def add_markdown_paragraph(doc, text):
    """Add a paragraph handling various inline formatting."""
    if not text.strip():
        doc.add_paragraph()
        return

    # Handle bold
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p

def parse_table_line(line):
    """Parse a markdown table row."""
    line = line.strip().strip('|')
    return [cell.strip() for cell in line.split('|')]

def add_table_from_markdown(doc, lines_iter):
    """Read a markdown table and add it to the doc."""
    # Collect table lines
    table_lines = []
    for line in lines_iter:
        if line.strip().startswith('|') and not line.strip().startswith('>'):
            table_lines.append(line.strip())
        else:
            # This line is not part of the table, put it back
            break

    if not table_lines:
        return

    # First line is header, second is separator
    if len(table_lines) < 2:
        return

    headers = parse_table_line(table_lines[0])
    # Skip separator line (---|---|---)
    data_rows = []
    for tl in table_lines[2:]:
        data_rows.append(parse_table_line(tl))

    if not data_rows:
        return

    table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_i, row_data in enumerate(data_rows):
        row_cells = table.rows[row_i + 1].cells
        for col_i, cell_text in enumerate(row_data):
            if col_i < len(headers):
                row_cells[col_i].text = cell_text
                for p in row_cells[col_i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table

# Parse and build the document
lines = article.split('\n')
i = 0

# Image insertion positions (based on T6_final_头条.md structure)
# We'll track sections and insert images at the right spots

images = {
    '封面图': 'T5_封面图.png',
    '价格梯度图': 'T5_价格梯度图.png',
    '品牌对比图': 'T5_品牌对比图.png',
    '选购指南图': 'T5_选购指南图.png',
    '618攻略图': 'T5_618攻略图.png',
}

image_inserted = {
    '封面': False,
    '价格梯度': False,
    '品牌对比': False,
    '选购指南': False,
    '618攻略': False,
}

image_captions = {
    'T5_封面图.png': '三大品牌——海尔云溪、小天鹅小乌梅、西门子小晶钻——2026年618选购全览',
    'T5_价格梯度图.png': '三大品牌全型号价格分布——海尔覆盖¥665-4,279，小天鹅覆盖¥669-2,779，西门子覆盖¥2,522-9,241',
    'T5_品牌对比图.png': '海尔/小天鹅/西门子八大维度横向对比——含上榜数量、洗净比、静音表现、售后覆盖、价格竞争力',
    'T5_选购指南图.png': '8种选购需求场景——千元入门至高端旗舰，各品牌对应推荐型号与价格',
    'T5_618攻略图.png': '2026年618洗衣机购买攻略——5/31开门红至6/15高潮期时间轴、国补叠加公式、购买清单',
}

in_table = False
table_lines = []
skipping_checklist = False

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip T3 quality checklist
    if stripped.startswith('## T3 质量门禁'):
        break
    if stripped == '---' and i > len(lines) - 10:
        break

    # Skip separator lines
    if stripped == '---':
        i += 1
        continue

    # Handle H1 title
    if stripped.startswith('# ') and not stripped.startswith('## '):
        title_text = stripped[2:]
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Insert cover image right after title
        img_path = os.path.join(BASE, images['封面图'])
        if os.path.exists(img_path) and not image_inserted['封面']:
            doc.add_paragraph()
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            # Caption
            cap = doc.add_paragraph(image_captions[images['封面图']])
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.style = doc.styles['Normal']
            for r in cap.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()
            image_inserted['封面'] = True

        i += 1
        continue

    # Handle H2 headings
    if stripped.startswith('## '):
        heading_text = stripped[3:]

        # Insert 价格梯度图 before "同价位硬碰硬" heading
        if '同价位硬碰硬' in heading_text and not image_inserted['价格梯度']:
            img_path = os.path.join(BASE, images['价格梯度图'])
            if os.path.exists(img_path):
                doc.add_paragraph()
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                cap = doc.add_paragraph(image_captions[images['价格梯度图']])
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                doc.add_paragraph()
                image_inserted['价格梯度'] = True

        doc.add_heading(heading_text, level=1)
        i += 1
        continue

    # Handle H3 headings
    if stripped.startswith('### '):
        heading_text = stripped[4:]
        doc.add_heading(heading_text, level=2)
        i += 1
        continue

    # Handle blockquotes (data source notes)
    if stripped.startswith('> '):
        quote_text = stripped[2:]
        p = doc.add_paragraph()
        run = p.add_run(quote_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.italic = True
        p.paragraph_format.left_indent = Cm(0.5)
        i += 1
        continue

    # Handle markdown tables
    if stripped.startswith('|') and not stripped.startswith('>'):
        table_lines = [stripped]
        j = i + 1
        while j < len(lines) and lines[j].strip().startswith('|'):
            table_lines.append(lines[j].strip())
            j += 1

        if len(table_lines) >= 3:  # header + separator + at least 1 data row
            headers = parse_table_line(table_lines[0])
            data_rows = [parse_table_line(tl) for tl in table_lines[2:]]

            table = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header
            for col_i, header in enumerate(headers):
                cell = table.rows[0].cells[col_i]
                cell.text = header
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(8)

            # Data
            for row_i, row_data in enumerate(data_rows):
                for col_i, cell_text in enumerate(row_data):
                    if col_i < len(headers):
                        cell = table.rows[row_i + 1].cells[col_i]
                        cell.text = cell_text
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(8)

            doc.add_paragraph()  # spacing

        i = j
        continue

    # Handle list items
    if stripped.startswith('- '):
        text = stripped[2:]
        p = doc.add_paragraph(style='List Bullet')
        # Handle bold markers in list
        parts = re.split(r'(\*\*.*?\*\*)', text)
        p.clear()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue

    # Regular paragraph
    if stripped:
        # Check image insertion points

        # 品牌对比图 - after the big comparison table paragraph
        if '数据来源：京东到手价' in stripped and not image_inserted['品牌对比']:
            # Add the paragraph first
            add_markdown_paragraph(doc, stripped)
            i += 1

            img_path = os.path.join(BASE, images['品牌对比图'])
            if os.path.exists(img_path):
                doc.add_paragraph()
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                cap = doc.add_paragraph(image_captions[images['品牌对比图']])
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                doc.add_paragraph()
                image_inserted['品牌对比'] = True
            continue

        # 选购指南图 - after "租房或预算紧张" paragraph
        if '租房或预算紧张' in stripped and not image_inserted['选购指南']:
            add_markdown_paragraph(doc, stripped)
            i += 1

            img_path = os.path.join(BASE, images['选购指南图'])
            if os.path.exists(img_path):
                doc.add_paragraph()
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(img_path, width=Inches(5.0))
                cap = doc.add_paragraph(image_captions[images['选购指南图']])
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                doc.add_paragraph()
                image_inserted['选购指南'] = True
            continue

        # 618攻略图 - after "别指望着618能看到比去年还低的绝对价格" paragraph
        if '别指望着618能看到比去年还低的绝对价格' in stripped and not image_inserted['618攻略']:
            add_markdown_paragraph(doc, stripped)
            i += 1

            img_path = os.path.join(BASE, images['618攻略图'])
            if os.path.exists(img_path):
                doc.add_paragraph()
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_p.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                cap = doc.add_paragraph(image_captions[images['618攻略图']])
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                doc.add_paragraph()
                image_inserted['618攻略'] = True
            continue

        add_markdown_paragraph(doc, stripped)
    else:
        doc.add_paragraph()

    i += 1

# Add footer note
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— END —')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

output_path = os.path.join(BASE, '2026年618洗衣机选购指南_终稿.docx')
doc.save(output_path)
print(f'Word document saved to: {output_path}')
print(f'Images embedded: {sum(1 for v in image_inserted.values() if v)}/5')
for k, v in image_inserted.items():
    print(f'  {k}: {"✅" if v else "❌"}')
