"""从T6_final_头条.md生成Word文档，嵌入T5配图"""
import re, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/aaron/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/ARTICLE WRITER/outputs"
MD_PATH = os.path.join(BASE, "T6_final_头条.md")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.8
style.paragraph_format.space_after = Pt(6)

with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')
    stripped = line.strip()

    # 空行跳过
    if not stripped:
        i += 1
        continue

    # 分隔线
    if stripped == '---':
        i += 1
        continue

    # 图片行
    img_match = re.match(r'!\[.*?\]\((.*?)\)', stripped)
    if img_match:
        img_file = os.path.join(BASE, img_match.group(1))
        if os.path.exists(img_file):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_file, width=Inches(5.8))
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        i += 1
        continue

    # 标题 H1
    if stripped.startswith('# ') and not stripped.startswith('## '):
        title_text = stripped[2:]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = '微软雅黑'
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_after = Pt(16)
        i += 1
        continue

    # 标题 H2
    if stripped.startswith('## '):
        h2_text = stripped[3:]
        p = doc.add_paragraph()
        run = p.add_run(h2_text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = '微软雅黑'
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
        i += 1
        continue

    # 表格
    if stripped.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        # 过滤分隔行
        data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|$', l)]
        if len(data_lines) >= 2:
            headers = [c.strip() for c in data_lines[0].split('|')[1:-1]]
            rows = []
            for dl in data_lines[1:]:
                rows.append([c.strip() for c in dl.split('|')[1:-1]])
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ci, h in enumerate(headers):
                cell = table.rows[0].cells[ci]
                cell.text = h
                for par in cell.paragraphs:
                    for r in par.runs:
                        r.bold = True
                        r.font.size = Pt(10)
                        r.font.name = '微软雅黑'
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = val
                    for par in cell.paragraphs:
                        for r in par.runs:
                            r.font.size = Pt(10)
                            r.font.name = '微软雅黑'
            doc.add_paragraph()
        continue

    # 话题标签行
    if stripped.startswith('#') and not stripped.startswith('##') and len(stripped) < 30:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(stripped)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'
        run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
        i += 1
        continue

    # 引用块（> 开头）
    if stripped.startswith('>'):
        quote_text = stripped.lstrip('> ').strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.right_indent = Cm(1.5)
        run = p.add_run(quote_text)
        run.italic = True
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        i += 1
        continue

    # 数据来源行（斜体）
    if stripped.startswith('*') and stripped.endswith('*') and '数据来源' in stripped:
        src_text = stripped.strip('*').strip()
        p = doc.add_paragraph()
        run = p.add_run(src_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = '微软雅黑'
        i += 1
        continue

    # 普通段落 - 处理加粗标记
    p = doc.add_paragraph()
    # 按 **...** 拆分
    parts = re.split(r'(\*\*.*?\*\*)', stripped)
    for part in parts:
        bold_match = re.match(r'\*\*(.*?)\*\*', part)
        if bold_match:
            run = p.add_run(bold_match.group(1))
            run.bold = True
            run.font.name = '微软雅黑'
            run.font.size = Pt(11)
        else:
            run = p.add_run(part)
            run.font.name = '微软雅黑'
            run.font.size = Pt(11)
    i += 1

output_path = os.path.join(BASE, 'AI最大变化不是变聪明了_头条图文.docx')
doc.save(output_path)
sz = os.path.getsize(output_path) / 1024 / 1024
print(f'Word文档已生成：{output_path}')
print(f'文件大小：{sz:.1f} MB')
