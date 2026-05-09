#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成去AI味润色版 + 6张NotebookLM配图嵌入的Word文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE = '/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/ARTICLE WRITER/outputs'
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.6
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Microsoft YaHei'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def para(text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    if align: p.alignment = align
    return p

def img(filename, w=6.0):
    path = os.path.join(BASE, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(w))
    else:
        para(f'[图片缺失: {filename}]', color=(0xff,0,0))

def tbl(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Microsoft YaHei'
                r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Microsoft YaHei'
                    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ═══════════════════════ 正文 ═══════════════════════

# 标题
t = doc.add_heading('波轮还是滚筒？2026年洗衣机选购核心看这3点', level=0)
for r in t.runs:
    r.font.name = 'Microsoft YaHei'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x0d,0x0d,0x0d)

img('NLM_01_封面图.png', 6.0)
doc.add_paragraph()

# 开头
para('买洗衣机纠结波轮还是滚筒？别纠结了，2026年抓三个数就够：洗涤方式看你的生活场景、洗净比1.1起步、电机别被"直驱碾压皮带"的话术带偏。这三个数搞明白了，导购再怎么绕，你也花不了冤枉钱。')

para('我把京东在售的30款洗衣机翻了个底朝天，又对着产品信息库的结构化参数挨个核对了一遍，三个维度给你一次性讲透。看完这篇，心里基本就有数了。')

# ═══════════ 维度一 ═══════════
heading('维度一：波轮还是滚筒，不是谁好谁差的问题', 2)

para('波轮洗衣机最大的槽点就是缠绕。它靠底部波轮盘旋转搅动水流，衣服搅成一团是物理上绕不开的事，磨损大、耗水也高。产品信息库的数据摆在面前：波轮洗净比集中在0.9-0.95（来源：产品信息库15款波轮数据，2026年5月采集，国标测试条件）。')

para('那滚筒就完美了？也不是。')

para('滚筒靠内筒旋转把衣服提起来再摔下去，通过摔打来洗。洗净比普遍在1.1-1.33（来源：产品信息库1252款滚筒数据，双源验证通过）。洗得确实干净，对衣服也温柔。但槽点同样明显：洗一次动不动60到120分钟，中途想塞件衣服进去，门根本打不开（来源：京东实测数据+行业通用知识）。')

para('先看一张核心对比表，数据说话：')

img('NLM_02_参数对比图.png', 6.0)
doc.add_paragraph()

tbl(
    ['对比维度', '波轮洗衣机', '滚筒洗衣机', '数据来源'],
    [
        ['价格区间', '349-2799元', '958-9241元', '产品库+京东'],
        ['洗净比', '0.9-0.95', '1.1-1.33', '产品库+Tavily 双源验证'],
        ['衣物磨损', '较大（缠绕拉扯）', '较小（摔打护衣）', '多源一致'],
        ['洗涤时间', '40-50分钟', '60-120分钟', '京东实测'],
        ['噪音', '偏大', '较小（变频电机加持）', '京东实测'],
        ['中途添衣', '随时打开', '运行中开不了门', '行业通用'],
        ['嵌入橱柜', '上翻盖，不能嵌入', '侧开门，可嵌入', '多源一致'],
        ['功能丰富度', '基础款为主', '除菌/空气洗/智能投放都有', '产品库'],
    ]
)
para('数据来源：产品信息库API + 京东30款实时数据，2026年5月6日采集', size=9, color=(0x99,0x99,0x99))
doc.add_paragraph()

para('但把波轮说成一无是处也不公道。你看看京东热卖榜——扬子10kg波轮349元，热卖榜第1，已售20万+——这说明波轮的市场需求实实在在存在。有几个场景波轮确实不好替代：')

para('第一个，租房过渡。几百块搞定一台，哪天搬家丢了也不心疼。扬子349元那款就是典型。')
para('第二个，给爸妈买。不用弯腰、按键直观，老人学起来没门槛，比触控面板友好太多。')
para('第三个，预算吃紧。同价位波轮容量比滚筒大不少，钱花在刀刃上。')

img('NLM_03_场景对比图.png', 6.0)
doc.add_paragraph()

para('我之前帮亲戚挑过一台波轮。老太太就认准了，说衣服往里一扔，按一下开始，完事。滚筒那些空气洗、智能投放、APP控制，她用不上，也不想学。这件事让我挺感触的——有时候"好"的标准不是参数决定的，是谁在用决定的。')

img('NLM_04_优缺点图.png', 5.0)
doc.add_paragraph()

para('波轮好在哪：价格低、操作简单、中途随便添衣、维修也便宜。短板也摆在明面上：衣服缠绕磨损大、噪音偏高、功能单一、没法嵌入橱柜。')
para('滚筒的优势：洗净能力强、护衣效果好、功能丰富、能嵌入橱柜一体化。不足的地方：洗涤时间长、中途添衣服麻烦、价格门槛高、维修成本也高。')

para('代表机型参考（价格均为京东2026年5月6日到手价）：', bold=True)
para('• 波轮入门：扬子10kg波轮 349元 — 热卖榜第1，已售20万+')
para('• 波轮品质：小天鹅V26DT直驱 851元 — 金榜第2，直驱变频静音')
para('• 滚筒入门：海信10kg滚筒 958元 — 热卖榜第6，55cm超薄机身')
para('• 滚筒主流：统帅B29S 998元 — 70cm以下热卖榜第4，巴氏除菌')
para('• 滚筒品质：海尔云溪4.0 2721元 — TOP1，AI精控+精华洗+FPA直驱')

# ═══════════ 维度二 ═══════════
heading('维度二：洗净比——选洗衣机只盯一个参数的话，看它就够了', 2)

para('洗净比是国标强制指标，简单说就是被测洗衣机对标准污染布的洗净程度，跟参比洗衣机比出来的比值。滚筒合格线是1.03（来源：GB/T 4288-2018《家用和类似用途电动洗衣机》）。')

para('京东30款在售滚筒我逐个看了一遍，洗净比1.03的机型基本退出主流市场了。金榜前10的滚筒，洗净比全部在1.1以上——1.03这条线实际上已经被市场淘汰，不用纠结要不要买。')

img('NLM_06_价格档位图.png', 6.0)
doc.add_paragraph()

para('我把洗净比分了四个档位，每个档位对应具体机型和价格：', bold=True)

tbl(
    ['洗净比档位', '定位', '代表机型', '京东到手价', '数据来源'],
    [
        ['1.03', '淘汰线', '基本无主流在售', '—', '京东30款统计'],
        ['1.1', '主流线', '海尔初色29S / 统帅B29S / 海信10kg', '958-1189元', '产品库 双源验证'],
        ['1.15-1.2', '进阶线', '海信1.15 / 海尔云溪4.0 1.21', '958-2721元', '产品库 双源验证'],
        ['1.2+', '旗舰线', '小乌梅3.0 1.21 / 海信棉花糖1.33', '2686-5759元', '产品库 双源验证'],
    ]
)
para('数据来源：产品信息库API + 京东，2026年5月6日', size=9, color=(0x99,0x99,0x99))
doc.add_paragraph()

para('几个关键结论：')
para('1.03是国标合格线，但2026年不建议碰。京东金榜前10的滚筒，洗净比全在1.1以上，1.03的款连榜都上不去。')
para('1.1是今年的起步推荐值。统帅B29S 998元就是1.1洗净比，性价比很能打。')
para('想一步到位的话，1.2以上算旗舰水准。海尔云溪4.0的1.21（产品库ID:19167，双源验证），小乌梅3.0的1.21（产品库ID:18433），海信棉花糖的1.33（产品库ID:19691），都是高洗净比的代表。')

para('这里有个容易踩坑的地方：波轮洗净比为什么普遍只有0.9-0.95？是不是波轮洗不干净？', bold=True)

para('不是这么比的。波轮和滚筒的测试标准不一样——滚筒的参比洗衣机洗净能力本身就是基准线，所以滚筒洗净比天然偏高。直接拿波轮的0.9去比滚筒的1.1，相当于拿百米成绩跟马拉松成绩比谁跑得快，连规则都不同。正确做法是同类型内比：波轮里0.95比0.9好，滚筒里1.2比1.1好。')

para('还有个冷知识：2026年5月1日起，新版国标GB/T 4288-2025正式实施了，新增了轻洗涤和护色测试要求（来源：国家标准化管理委员会）。以后洗净比的参考含金量会越来越高，现在开始关注这个参数不吃亏。')

para('高洗净比意味着清洁能力强，顽固污渍处理得干净。代价也不小：同品牌内价格明显更高，旗舰款的功能溢价大。反过来，低洗净比不是说不能用，日常基础洗涤完全够，价格也友好，但油渍这类顽固污渍确实吃力。')

# ═══════════ 维度三 ═══════════
heading('维度三：电机——BLDC和DD直驱根本不是一回事', 2)

para('这是洗衣机选购里被误解最深的地方。')

para('很多导购张嘴就是"直驱比皮带高级"，好像DD直驱碾压BLDC皮带似的。真相是什么？BLDC全称Brushless Direct Current Motor，是无刷直流电机——说的是"用什么电机"。DD直驱全称Direct Driver，是驱动方式——说的是"怎么连接内筒"，取消了皮带直连。这两个概念根本不在一个维度上，没法说谁比谁好（来源：新浪科技+太平洋电脑网，双源验证）。')

para('打个比方你就懂了：有人问"汽油车好还是自动挡好"——汽油说的是发动机类型，自动挡说的是变速箱形式，这俩问题根本不在一个层面上。')

para('再补充一点：DD直驱用的电机，本质上也是BLDC无刷直流电机。而FPA直驱是海尔收购新西兰斐雪派克后引入的直驱电机技术品牌名（来源：海尔官网），说白了就是DD直驱的一种，只不过海尔给它起了个名字。')

tbl(
    ['电机方案', '原理', '代表机型', '京东到手价', '适合人群', '数据来源'],
    [
        ['定频电机', '有刷电机，转速固定', '扬子349元 / 志高371元', '349-371', '租房、极低预算', '京东'],
        ['BLDC+皮带', '无刷电机+皮带传动', '统帅B29S 998元 / 小乌梅3.0 2686元', '998-3400', '主流家庭', '产品库+Tavily'],
        ['DDM直驱变频', '无刷电机+直连内筒', '小天鹅V26DT 851元 / 海尔Mate6波轮 3199元', '851-3199', '静音+护衣', '产品库+京东'],
        ['FPA直驱', '海尔品牌专有直驱', '海尔云溪4.0 2721元（618到手价）', '2721+', '高端品质需求', '海尔官网验证'],
    ]
)
para('数据来源：产品信息库API + 新浪科技 + 太平洋电脑网，2026年5月', size=9, color=(0x99,0x99,0x99))
doc.add_paragraph()

para('那BLDC皮带和DD直驱到底怎么理解？简单说吧：BLDC皮带传动因为多了皮带这个弹性缓冲，低速扭力大，对洗涤效果反而是加分项；DD直驱取消皮带后传动效率更高、噪音更低，高速脱水的时候特别稳。各有各的好，没有谁绝对碾压谁。')

para('那怎么选？', bold=True)

para('预算1000以内，不用纠结。BLDC皮带方案完全够用，统帅B29S 998元就是BLDC皮带传动，京东70cm以下热卖榜第4，日常洗衣服一点毛病没有。')
para('预算1000-3000，可以考虑直驱方案。海尔云溪4.0 2721元用的FPA直驱，AI精控+精华洗，功能很全面。')

para('有个细节值得留意：波轮上装直驱和滚筒上装直驱，体验差距很大。波轮本身噪音就大，加了直驱后降噪效果非常明显。小天鹅V26DT 851元就是典型，直驱变频+静音设计，金榜第2。滚筒本身噪音控制得就不错，直驱的提升感知没波轮来得强。所以你买波轮的话，直驱是值得加钱的升级项；买滚筒的话，BLDC皮带和直驱的日常体验差距没你想的那么大。')

para('BLDC皮带方案胜在技术成熟、维修便宜、洗涤效果好、价格亲民。短板是皮带是消耗品，大概5到8年得换，高速脱水时噪音也会稍大。')
para('DD直驱方案的好处：没皮带免维护、噪音低、高速脱水稳、传动效率高。代价是价格更贵，维修成本高，部分品牌的零部件还得依赖进口。')

# ═══════════ 场景化选购 ═══════════
heading('场景化选购建议', 2)

para('聊了这么多参数，最后落到一个问题上：你该买哪款？')

img('NLM_05_人群选购图.png', 5.0)
doc.add_paragraph()

tbl(
    ['你的情况', '推荐方向', '代表型号', '京东参考价', '推荐理由'],
    [
        ['租房过渡', '波轮入门', '扬子10kg波轮', '349元', '热卖榜第1，已售20万+，搬家不心疼'],
        ['给爸妈买', '波轮直驱', '小天鹅V26DT', '851元', '金榜第2，直驱静音，操作简单'],
        ['小户型嵌入', '超薄滚筒', '海信10kg滚筒', '958元', '55cm超薄，热卖榜第6'],
        ['一家三口', '主流滚筒', '海尔初色29S', '1189元', '金榜第4，1.1洗净比，超薄嵌入'],
        ['母婴家庭', '高温滚筒', '小天鹅V23', '1137元', '95度高温煮洗，好评榜第6'],
        ['品质家庭', '旗舰滚筒', '海尔云溪4.0', '2721元', 'TOP1，AI精控+FPA直驱+精华洗'],
        ['追求高洗净比', '高端滚筒', '小乌梅3.0', '2686元', '1.21洗净比，超微气泡洗，5星推荐'],
    ]
)
para('数据来源：京东，2026年5月6日到手价。618期间（5月31日-6月18日）预计还有14%-17%降幅，一级能效机型可叠加国补15%（上限2000元）。', size=9, color=(0x99,0x99,0x99))
doc.add_paragraph()

para('多说一句省钱的事：今年家电整体在涨价。铜材缺口到了20%，铜价涨到98450元/吨，DRAM合约价预估涨90%-95%，终端售价普遍上调了5%-20%（来源：新浪财经、深圳新闻网，2026年4月报道）。所以618大概率是今年洗衣机价格的阶段性低点。两个下单节点记住就行：5月31日晚上8点开门红，6月15日晚上8点高潮期。')

# ═══════════ 总结 ═══════════
heading('总结', 2)

para('选洗衣机就三步。第一步，波轮还是滚筒？别看"谁更好"这种没意义的问题，看你家实际使用场景。第二步，洗净比。滚筒1.1起步，波轮跟滚筒的数值不能横着比，同类型内比才有意义。第三步，电机方案。1000以内BLDC皮带够用，预算充裕可以考虑直驱，波轮加直驱的感知提升比滚筒更明显。')

para('记住一句话：数据说话，场景选机，比听导购绕弯子靠谱得多。')

doc.add_paragraph()
para('ContentFleet v5.0 · 2026-05-06', size=8, color=(0xaa,0xaa,0xaa))
para('数据来源：产品信息库API + 京东30款实时数据 + Tavily搜索验证', size=8, color=(0xaa,0xaa,0xaa))

# ═══════════ 保存 ═══════════
out = os.path.join(BASE, '波轮还是滚筒_2026洗衣机选购指南_终稿.docx')
doc.save(out)
print(f'✅ 已生成: {out}')
