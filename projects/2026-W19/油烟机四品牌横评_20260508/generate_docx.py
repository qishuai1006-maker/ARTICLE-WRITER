#!/usr/bin/env python3
"""生成油烟机四品牌横评 · 带配图的 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUTS = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# 标题样式
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = '微软雅黑'
    h.font.bold = True
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)


def add_image(img_name, width_inches=5.5, caption=None):
    """插入图片"""
    path = os.path.join(OUTPUTS, img_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.style = doc.styles['Normal']
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()  # spacer


def add_body(text):
    """添加正文段落"""
    p = doc.add_paragraph(text)
    # 处理加粗标记
    for run in p.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
    return p


# ═══════════════════════════════════════
#  正文开始
# ═══════════════════════════════════════

# ── 标题 ──
title = doc.add_heading('油烟机推荐！2026四品牌横评不踩坑', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 封面图 ──
add_image('T5_封面图.png', 5.5, '▲ 2026油烟机四品牌横评（数据来源：京东商品页+品牌官方参数，2026年5月）')

# ── 开头 ──
add_body('2026年618马上到了，油烟机品牌扎堆推新，但我跟你说个实话：今年的选购逻辑已经变了。')

add_body('别一上来就比谁家风量大。30m³/min往上在大多数家庭里根本体验不出差别。真正拉开差距的，是「自清洁技术路线」。说白了，你现在买油烟机，选的不是哪个牌子，选的是这台机器三年后还能不能吸得动。')

add_body('我去年帮朋友挑油烟机，盯着参数表看了一周，最后发现最该关注的不是风量数字，而是自清洁方式。后来自己装的时候学乖了——先想清楚家里做什么菜、厨房什么结构，再对技术路线入座，路就清晰多了。')

# ── 核心参数扫盲 ──
doc.add_heading('先搞懂三个参数，别被数字绕晕', level=2)

add_body('油烟机的核心参数就三个，搞懂了就不会被导购带着走。')

p = doc.add_paragraph()
r = p.add_run('风量：管吸力，但别迷信数字')
r.bold = True
r.font.name = '微软雅黑'
add_body('风量的单位是m³/min，意思是每分钟能吸走多少立方米的空气。据京东商品页和品牌官方数据，目前主流机型在23-35m³/min之间。')
add_body('23m³/min够清淡饮食和日常家常菜。28m³/min是爆炒的及格线。到了30m³/min以上，说实话大多数人感知不到和28m³/min的差别，毕竟厨房不是密封实验室，开窗通风、空气对流都会影响实际效果。还有一个容易被忽略的点：排烟管多拐几个弯，实际风量衰减很明显。装修时烟管三四个弯头走下来，30m³/min的机器到手可能只剩20出头。所以走管的时候尽量走直线，少拐弯。')
add_body('结论很简单：28-30m³/min对绝大多数中国家庭来说绰绰有余，别为多出来的几个数字多花钱。')

p = doc.add_paragraph()
r = p.add_run('静压：管排力，高层住户首要关注的参数')
r.bold = True
r.font.name = '微软雅黑'
add_body('静压说的是油烟机往外排烟时对抗公共烟道阻力的能力，单位是Pa（帕斯卡）。')
add_body('住高层住宅的朋友注意了：公共烟道在饭点（晚上6-8点）是所有楼层同时往里排烟的，你住的楼层越低，面对的烟道压力越大。据ZOL数据，静压低于800Pa的机型在高层低户基本饭点必倒灌，你家做红烧肉邻居闻到无所谓，但你闻到别人家油烟，那滋味不好受。')
add_body('独栋、别墅、直接排外墙的房子，300-500Pa就够了。但如果你住10层以上的高层住宅、楼层又在1-10层，建议静压1000Pa起步，1200Pa以上才稳妥。这个参数比风量大5个单位重要得多。')

p = doc.add_paragraph()
r = p.add_run('自清洁：管寿命，决定你花四五千买的机器三年后还有没有用')
r.bold = True
r.font.name = '微软雅黑'
add_body('这是一个被严重低估的参数。据中家院检测数据，烟机内部油垢堆积后，风量衰减可达40%以上。翻译一下就是：4000块买的32m³/min油烟机，如果自清洁没做好，两年后实际吸力可能只相当于一台1500块的机器。')
add_body('市面上自清洁技术分四条路线：高温蒸汽洗、冲浪蒸水洗、免拆洗涂层、高温溶油洗。下一节会详细拆解这四条路线的差别，但你可以先记住一个判断：主动清洁（蒸汽洗/蒸水洗）的长期效果，明显好于被动防护（涂层/热干洗）。')

# 插入核心参数扫盲图
add_image('T5_核心参数扫盲图.png', 5.5, '▲ 油烟机三大核心参数一图看懂（图：NotebookLM）')

# ── 四大品牌主力机型 ──
doc.add_heading('四大品牌2026主力机型：各有各的技术护城河', level=2)

add_body('搞懂了参数，来看具体产品。我不跟你讲「哪个品牌好」，我跟你讲「什么情况该选谁」。')

# 美的
doc.add_heading('美的：蒸汽洗路线，省心派首选', level=3)
add_body('美的在油烟机领域的核心护城河就是蒸汽洗技术。110℃高温蒸汽软化油垢，高压水流冲刷，然后高速甩干，一个完整的主动清洁过程。')

p = doc.add_paragraph()
r = p.add_run('小西梅AK10')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约2800-3200元）：2026年主力新品，核心升级是「自动蒸汽洗2.0」。跟上一代最大的区别：隐藏式水盒（一次加水能用多次）+ 累计运行7小时自动触发清洗。你不用记着什么时候该洗，机器自己判断。参数方面，综合京东和品牌官方数据，风量30m³/min，最大静压1350Pa，机身厚度350mm能跟标准橱柜齐平。京东销量2000+。')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run('小西梅AK11 Pro')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（京东参考价约5400-5900元，国补后约4600-5000元）：AK10的升级版，据京东和品牌官方数据，风量32m³/min，最大静压1500Pa，加了MAX双风道。京东销量1万+，好评率99%。开放式厨房或者天天爆炒的家庭可以考虑，这个参数配置在5000元档属于第一梯队。')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run('AK7 Pro飓风版')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约2811元）：美的的销量王，京东已售10万+。据京东商品页数据，30m³/min风量，1250Pa静压，110℃高温蒸汽洗（一代，手动启动）。据中家院认证数据，洗净率99.1%。如果预算控制在3000以内，这款是经过大规模市场验证的选择。短板是蒸汽洗需要手动触发，没有AK10的自动识别功能。')
r.font.name = '微软雅黑'

add_body('美的的核心价值说白了就一条：用蒸汽洗对抗吸力衰减。110℃蒸汽加高压冲刷，是目前主动清洁里清洁力很突出的方案，适合不想每年花300块请人拆洗的家庭。')

# 老板
doc.add_heading('老板：免拆洗+AI路线，性能派首选', level=3)
add_body('老板走的是另一条技术路线。核心技术叫「油立净涂层」，在烟机内壁涂一层纳米疏油材料，让油污不容易附着，再配合高油脂分离率把油拦截在外。思路是「不让油沾上」，而不是「沾上了再洗」。')

p = doc.add_paragraph()
r = p.add_run('小环翼C5')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约4000元以上）：2026年全新旗舰。据品牌官方数据，风量32m³/min，最大静压1700Pa，参数配置在四品牌中处于领先位置。搭载AI大模型和鲲鹏风道（竖置涡轮），噪音低至42dB（AI全域降噪），6吸烟口设计。C5的竖置涡轮是个有意思的设计：传统烟机蜗壳横置，气流拐弯多；竖置涡轮让气流直进直出，排烟效率更高。京东作为新品销量还没起来，价格也是四款里最高的。')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run('双子星E1')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约2986-3200元）：据京东数据，28m³/min风量，1450Pa静压，搭载灵犀烟灶联动3.0和AI九档调节，京东已售8万+。3000预算想买老板的，E1比D1P多了450Pa静压，高层住户更合适。')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run('小黑翼D1P')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约2573元）：老板的性价比担当，京东已售10万+。据京东商品页数据，27m³/min风量，1000Pa静压，免拆洗涂层。有个实用设计：蒸煮模式烟板自动收起，蒸鱼煮粥时不会把蒸汽全吸走造成浪费。')
r.font.name = '微软雅黑'

add_body('老板的核心价值就一条：极致排烟性能加不依赖水洗。C5的1700Pa是四品牌中静压参数最大的，高层低户怕倒灌的可以重点看。')

# 方太
doc.add_heading('方太：拢烟+智感净味路线，品质派首选', level=3)
add_body('方太的思路跟前两家不一样。核心不在自清洁，而在「油烟全过程管理」——不是等油烟出来了再吸，而是从产生到排出全过程管控。')

p = doc.add_paragraph()
r = p.add_run('V10/V10-G')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约4500-5139元）：30m³/min风量，0缝平嵌设计，支持APP智控。核心功能是「智感净味」，24小时空气质量哨兵模式，待机状态下也在监测厨房空气质量，自动启动排风。京东已售2万+（套装口径）。价格是四款里最贵的，但方太的品控和售后在行业里口碑确实不错。')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run('V1S/V1S-G')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约4000-4594元）：据京东和品牌官方数据，28m³/min风量，1450Pa最大静压，0缝真平嵌，磁吸油网。京东已售4万+。在意厨房颜值的话，平嵌设计确实比传统油烟机凸出橱柜一大截好看很多。')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run('F5')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约2500-3483元）：27m³/min风量，1300Pa静压，35寸超大云板拢烟。据京东商品页数据，这是3000元预算里方太很值得考虑的型号。35寸云板意味着拢烟范围大，颠勺、爆炒这类向上油烟能兜得住。不足是传统欧式顶吸形态，身高175cm以上的人可能有碰头风险。')
r.font.name = '微软雅黑'

add_body('方太的核心价值是颜值加体验加服务。自清洁用的是高温溶油洗，电机70℃发热软化油垢后离心甩出，清洁力不如蒸汽洗，但对于轻油烹饪的家庭完全够用。')

# 华帝
doc.add_heading('华帝：蒸水洗+颜值路线，参数党首选', level=3)
add_body('华帝是四家里自清洁技术路线最为多样的品牌。核心技术是「冲浪蒸水洗」，先高温蒸汽软化油垢，再用脉冲高压热水冲刷，两步走。')

p = doc.add_paragraph()
r = p.add_run('小飞碟Max S66')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约2597元）：28m³/min风量，1200Pa最大静压，第四代冲浪自动洗，AG玻璃面板加0缝齐平设计，京东已售2万+。S66在2000-3000元档是很有竞争力的选手。AG玻璃面板防指纹不显脏，0缝齐平不积灰，颜值在同价位里表现突出。据品牌官方数据，洗净覆盖率超99%。')
r.font.name = '微软雅黑'

p = doc.add_paragraph()
r = p.add_run('隐系列Y9')
r.bold = True
r.font.name = '微软雅黑'
r = p.add_run('（约4000元以上）：2026年旗舰，据品牌官方数据，32m³/min风量，1600Pa最大静压，第五代蒸水洗。洗净率99.2%，灭菌率99.99%，在行业公开数据中处于领先水平。获得AWE2026艾普兰创新奖。105mm超薄机身，24小时VOC哨兵功能。预算到4000以上、对自清洁有极致追求的，Y9值得认真考虑。')
r.font.name = '微软雅黑'

add_body('华帝的核心价值是高洗净率加出色颜值。S66是2000-3000档很值得关注的型号之一，Y9的洗净率数据是当前可查数据中的领先水平。')

# 插入参数对比图
add_image('T5_参数对比图.png', 5.5, '▲ 四大品牌2026主力机型核心参数对比（数据来源：京东商品页+品牌官网，2026年5月）')

# ── 自清洁四条技术路线 ──
doc.add_heading('自清洁四条技术路线，到底怎么选？', level=2)

add_body('这是全文最核心的部分。前面说了，自清洁决定油烟机三年后还能不能干活。')

# 表格
table = doc.add_table(rows=9, cols=5, style='Light Shading Accent 1')
# 表头
headers = ['维度', '高温蒸汽洗（美的）', '冲浪蒸水洗（华帝）', '免拆洗涂层（老板）', '高温溶油洗（方太）']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.name = '微软雅黑'

data = [
    ['原理', '110℃蒸汽软化+高压冲刷', '蒸汽软化+脉冲热水冲刷', '纳米疏油涂层防附着', '电机70℃发热+离心甩油'],
    ['清洁方式', '主动清洁', '主动清洁', '被动防护', '被动防护'],
    ['洗净率', '99.1%（中家院认证）', '99.2%（品牌官方数据）', '无法量化', '无法量化'],
    ['需加水', '是（2.0版隐藏水盒）', '是（内置水箱）', '否', '否'],
    ['自动化', '蒸汽洗2.0自动触发', '60h提醒/部分自动', '零操作', '零操作'],
    ['长期衰减', '低', '最低', '中（涂层有效期3-5年）', '中高（重油清洁力有限）'],
    ['使用成本', '零（自动触发）', '零', '3-5年后可能需拆洗', '最终可能需拆洗'],
    ['适合烹饪', '爆炒/煎炸/重油', '爆炒/煎炸/重油', '中等油烟', '轻油/蒸煮'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.name = '微软雅黑'

doc.add_paragraph()

add_body('看懂这个表，你就理解为什么我说「选油烟机其实是选自清洁技术路线」了。')
add_body('蒸汽洗和蒸水洗是真用水和蒸汽去冲刷内部油垢，主动出击。免拆洗涂层和热干洗是设防线，被动防御。据多个长期使用评测数据，主动清洁方案在三年后的吸力保持率明显优于被动防护方案。')
add_body('不过公平地说，被动防护有个好处：零操作、零维护。如果你是轻油烹饪，平时蒸个鱼煮个粥，被动防护完全够用。但你家如果是川湘爆炒风格，建议还是选主动清洁路线，别跟自己的使用习惯较劲。')

# 插入自清洁技术对比图
add_image('T5_自清洁技术对比图.png', 5.0, '▲ 油烟机自清洁四大技术路线深度对比（图：NotebookLM）')

# ── 场景选购 ──
doc.add_heading('按你家的情况，直接对号入座', level=2)

add_body('不讲大道理了，直接给购买建议。以下价格均为京东单品参考价，截至2026年5月。')

advice = [
    ('开放式厨房+天天爆炒', '美的小西梅AK11 Pro（32风量/1500Pa/蒸汽洗2.0自动触发）。AK11 Pro在旗舰价位段把风量、静压、自清洁三者拉满了。预算有限可以看AK10（30风量/1350Pa/约2800-3200元），蒸汽洗2.0核心功能一样不少。'),
    ('高层低楼层（1-10楼/总高20层+）', '老板C5（1700Pa，领先的静压表现）或华帝Y9（1600Pa）。静压是你首要关注的参数，别的参数都得给它让路。'),
    ('预算2500-3000元', '美的AK7 Pro（约2811元/10万+销量）或华帝S66（约2597元/AG玻璃平嵌/2万+销量）或老板D1P（约2573元/10万+销量）。三款都是经过大规模市场验证的型号，销量和口碑能说明问题。'),
    ('在意颜值+不想拆洗', '华帝S66（AG玻璃+0缝齐平+自动冲浪洗）或美的AK10（350mm超薄+蒸汽洗2.0自动触发）。这两款在颜值和实用之间平衡得不错。'),
    ('清蒸为主+追求安静', '老板C5（42dB AI降噪）或方太F5（35寸大云板）。风量不用太大，但拢烟能力和噪音控制会更影响日常体验。'),
]
for title_text, detail in advice:
    p = doc.add_paragraph()
    r = p.add_run(f'▸ {title_text}：')
    r.bold = True
    r.font.name = '微软雅黑'
    r = p.add_run(detail)
    r.font.name = '微软雅黑'

# 插入场景选购图
add_image('T5_场景选购图.png', 5.5, '▲ 按厨房场景速查推荐——2026油烟机选购地图（图：NotebookLM）')

# ── 新老款对比 ──
doc.add_heading('别花新款的钱买到去年的库存', level=2)

add_body('2026年油烟机市场有个现象：品牌扎堆推新，但不是所有新品都值得多花钱。')

add_body('真正实现代际升级的：美的小西梅系列（蒸汽洗2.0从手动变自动+隐藏水盒）、老板C5（竖置涡轮+AI大模型+1700Pa静压）。这些是真正的技术跃迁，多花的钱能换来不一样的体验。')

add_body('换壳为主的：有些品牌2026新品跟2025款对比，核心风机和参数几乎一样，面板颜色换了、名字改了就当新品卖。据拆机实测数据，线上款和线下款的核心风机完全一样，所以别被「新款」两个字迷惑，看参数比看年份靠谱。')

add_body('老款什么时候买：AK7 Pro、D1P、F5这些10万+销量验证的口碑款，618叠加国补15%降价后性价比更高。建议盯住京东实时价格，新款和老款价差超过500元时，老款的性价比优势就很明显。京东搜索「领国补855」或「家电国补687」可查看当前补贴入口。')

# 插入新老款对比图
add_image('T5_新老款对比图.png', 5.5, '▲ 2026年油烟机新老款交替对照——谁是真换代谁是换壳（图：NotebookLM）')

# ── 结尾 ──
doc.add_heading('最后说两句', level=2)

add_body('2026年的油烟机竞争，已经从「谁吸力大」变成了「谁吸力保持得久」。风量30m³/min的时代已经到来且正在过剩，真正拉开长期体验差距的，是自清洁技术路线。')

add_body('选油烟机记住三件事：静压管排力，自清洁管寿命，风量够用就行。618叠加国补15%是换新的好时机，但别急着下单——先想清楚你家厨房是什么条件、做什么菜，再对技术路线入座。')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— END —')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.font.name = '微软雅黑'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('数据来源：京东商品页 + 品牌官方参数 | 采集时间：2026年5月 | ContentFleet v6.1')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
r.font.name = '微软雅黑'

# ── 保存 ──
output_path = os.path.join(OUTPUTS, '油烟机四品牌横评_2026.docx')
doc.save(output_path)
print(f'✅ Word 文档已生成: {output_path}')
