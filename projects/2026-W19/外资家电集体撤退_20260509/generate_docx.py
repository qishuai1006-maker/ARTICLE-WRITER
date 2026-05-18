#!/usr/bin/env python3
"""生成 外资家电集体撤退 头条深度图文 Word 文档"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUTS = os.path.join(BASE_DIR, "outputs")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Microsoft YaHei'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.line_spacing = 1.8
style_normal.paragraph_format.space_after = Pt(6)
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


def add_source(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_image(filename, width_inches=6.0):
    path = os.path.join(OUTPUTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
    else:
        add_para(f"[图片缺失: {filename}]")


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Microsoft YaHei'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.add_paragraph()


# ════════════════════════════════════════════
#  ARTICLE CONTENT
# ════════════════════════════════════════════

add_title("外资家电集体撤退：三星/索尼/LG走了，谁在填满中国人的家？")

add_image("T5_封面图.png", 5.2)

add_para("5月6日，三星电子在官网发了一条不到200字的公告：即日起在中国大陆停售所有家电产品。电视、冰箱、洗衣机、空调、吸尘器——13个品类，一刀切。")

add_para("但三星不是第一个走的，甚至不是今年第一个。把过去十二年退出或大幅收缩中国市场的外资家电品牌列一张名单，数量远比想象中多：三洋、东芝、飞利浦、夏普、惠而浦、LG、索尼、松下（电视业务）、A.O.史密斯、博西家电，加上三星，11个品牌。")

add_para("一个时代在收尾。真正值得关心的问题不是“它们为什么走了”，而是“走了之后，中国家庭里摆着的电器是谁在造”。")

# ── Section 1 ──
add_h2("一张撤退名单：12年，11个品牌")

add_para("先看全貌。")

add_table(
    ["品牌", "国别", "退出/收缩方式", "时间"],
    [
        ["三洋", "日本", "被惠而浦收购，品牌全线退出中国", "2014年"],
        ["夏普", "日本", "被富士康收购后持续收缩，已边缘化", "2016年起"],
        ["东芝", "日本", "电视授权海信、白电售予美的、关闭大连工厂", "2018-2021年"],
        ["惠而浦", "美国", "中国业务被格兰仕收购", "2021年"],
        ["飞利浦", "荷兰", "44亿美元将家电业务卖给高瓴资本", "2020年"],
        ["LG", "韩国", "全线撤出中国大陆，N+5倍赔偿，业务转越南", "2024年底"],
        ["A.O.史密斯", "美国", "在华销售额下滑12%，传闻海信接触收购", "2025年"],
        ["博西家电", "德国", "大中华区营收下滑7.1%，全球范围内下滑明显的市场", "2025年"],
        ["松下", "日本", "关停电视工厂，电视交创维运营", "2025年"],
        ["索尼", "日本", "电视业务托管给TCL（TCL持股51%合资运营）", "2025年"],
        ["三星", "韩国", "全线停售13个家电品类，手机不受影响", "2026年5月"],
    ]
)
add_source("数据来源：各品牌公开公告及行业研报综合整理")

add_image("T5_外资撤退全景.png")

add_para("11个品牌，横跨日韩欧美，几乎涵盖了中国消费者能叫出名字的所有外资家电。退出方式也值得细看：东芝把电视给了海信、白电给了美的；飞利浦把整个家电部门44亿美元打包卖给了高瓴；索尼更直接——跟TCL合资公司，自己只出品牌，TCL出技术和产线。")

add_para("说白了，这不是“战略调整”，是交了钥匙。")

add_para("两组数据能说明撤退的规模。电视市场：2025年，中国八大国产电视品牌（海信、TCL、小米、创维、长虹、海尔、康佳、华为）合计出货3096.3万台，占94.1%。所有外资品牌加在一起，出货量不到100万台，份额不到5%。")
add_source("数据来源：洛图科技（RUNTO）《中国电视市场品牌出货月度追踪》")

add_para("白电差距更大。冰箱领域，海尔一家独占47.7%——将近半壁江山。加上海信（16.4%）和美的（14.8%），前三家拿走了将近80%。外资品牌（博西、松下）合计份额不到15%，而且年年在缩。")
add_source("数据来源：奥维云网（AVC）2025年监测数据")

# ── Section 2 ──
add_h2("分品类：谁在接管中国家庭的客厅和厨房")

add_image("T5_品类替代格局.png")

add_para("外资品牌撤退留下的空白，不是一夜之间被填上的。每个品类都有自己的竞争故事。")

add_image("T5_客厅电视.jpg", 5.5)

add_h3("电视：海信和TCL在三星的主场打赢了它")

add_para("这是替代最彻底的品类。2025年全球电视出货量：三星3530万台（连续20年全球出货量居首），TCL 3040万台，海信2990万台。三星和TCL的差距从曾经的遥不可及，缩小到了490万台。2025年12月，TCL单月出货量首次超过三星，登顶全球月度榜首。")
add_source("数据来源：奥维睿沃（AVC Revo）2025年全球彩电市场报告")

add_para("技术层面的弯道超车是关键。TCL连续五年全球Mini LED电视出货量第一，2025年上半年Mini LED出货137万台，同比暴增176%。海信在百吋以上大屏连续三年全球出货量榜首，2025年份额57.1%，激光电视全球七连冠，市占率70.3%。")

add_para("价格差距更致命。同规格Mini LED产品，国产品牌普遍比三星便宜三到五成。海信Mini LED入门款4599元，三星同技术最低也要11000元以上。0到4000元这个中国家庭最集中的价格区间，三星几乎没有产品覆盖。")
add_source("数据来源：奥维云网（AVC）及品牌公开数据")

add_h3("冰箱：海尔一家吃掉近半壁江山")

add_para("冰箱市场格局最集中。海尔线下份额47.7%，加上海信（16.4%）和美的（14.8%），国产品牌前三就吃掉了将近80%。")

add_image("T5_现代厨房冰箱.jpg", 5.5)

add_para("高端市场的变化更值得关注。1.5万以上的高端冰箱市场，海尔旗下的卡萨帝份额超过71%。这个价位曾经是西门子、博世、松下的自留地。卡萨帝的策略很清晰——以最高均价拿最高份额，不是靠便宜，是在同一个价位段正面赢了。")

add_table(
    ["价位段", "海尔系份额", "外资最高份额（博西）"],
    [
        ["8000-10000元", "47.3%", "9.7%"],
        ["15000-20000元", "60.8%", "18%"],
    ]
)
add_source("数据来源：GfK中怡康2025-2026年监测数据")

add_h3("洗衣机：海尔加小天鹅，两家拿走七成")

add_para("洗衣机市场集中度更高。海尔线下份额47.3%，美的旗下的小天鹅紧随其后，两家加起来超过70%。外资品牌（博西5-8%、松下3-5%）已经退出了主流竞争。")

add_para("高端洗衣机的替代同样迅速。卡萨帝在1万以上市场份额达到76.3%。洗烘一体2万+价位段，卡萨帝份额突破95%。如果两年前花2万买了一台西门子洗烘一体，今天同价位的卡萨帝在AI智投、精华洗、主动服务等功能上已经拉开了代差。")

add_image("T5_洗衣机家居.jpg", 5.5)

add_source("数据来源：GfK中怡康2026年1-14周监测数据")

add_h3("空调：美格力海尔三强，小米从侧翼杀入")

add_para("空调市场格局相对稳定：美的约29%排第一，格力约17%（份额同比下滑2个百分点），海尔约15%。但2025年的变量是小米——线上5月份额一度达到18%，单月超过格力。")

add_para("小米空调2025年出货超850万台，同比增长24%。背后的逻辑不是参数明显领先，而是生态绑定：当一个家庭已经有小米电视、小爱音箱、扫地机器人，下一台空调大概率还是小米，因为“一句话联动”的体验太顺滑了。AIoT平台连接设备10.79亿台，米家App月活1.127亿——这种生态粘性，靠卖单一硬件的传统品牌很难复制。")
add_source("数据来源：小米集团2025年财报")

add_h3("厨电：外资从来没站稳过")

add_para("油烟机、灶具、集成灶——这个品类外资品牌始终没有建立起竞争力。方太、老板、美的、华帝四大国产品牌占据了绝大部分市场份额。2026年的厨电市场，找外资品牌反而是件难事。")

# ── Section 3 ──
add_h2("还没走的，凭什么留下")

add_para("不过，不是所有外资品牌都在跑。松下在2026年反而在加码中国——完成集团重组后，决策权下放到“中国东北亚事业”板块，口号是“中国的事情在中国决定”。削减低毛利产品，聚焦“住空间解决方案”，直营品类销售同比翻倍。")
add_source("数据来源：新浪财经报道")

add_para("松下能留下，是因为做对了两件事：决策权本地化，不再事事请示大阪总部；找到了国产品牌暂时覆盖不够深的差异化赛道——纳诺怡吹风机、全屋净水解决方案。")

add_para("除了松下，还有几个细分领域外资品牌仍有优势：")

add_para("大金：中央空调和VRV系统，技术和工程能力仍是行业标杆")
add_para("戴森：吹风机HD15至今没有真正的国产平替，吸尘器V系列同样领先")
add_para("博西家电：高端嵌入式洗碗机和洗烘一体机仍有竞争力")
add_para("A.O.史密斯：高端热水器和商用净水系统")

add_para("这些品牌能留下，靠的不是品牌光环，而是在具体品类上有真正难以替代的技术壁垒。")

# ── Section 4 ──
add_h2("2026年买家电，记住这几张表")

add_image("T5_2026选购指南.png")

add_h3("分品类国产替代首选")

add_table(
    ["品类", "首选", "备选", "参考价位", "关键看什么"],
    [
        ["电视", "海信、TCL", "小米", "4000-8000元", "Mini LED分区数、HDR峰值亮度、低反膜"],
        ["冰箱", "海尔/卡萨帝", "容声、美的", "3000-15000元", "双系统双循环、净味技术、容积率"],
        ["洗衣机", "海尔/卡萨帝", "小天鹅、美的", "2000-8000元", "直驱电机、精华洗、洗烘一体"],
        ["空调", "美的", "格力、海尔、小米", "2000-4000元", "APF能效值、实际使用体验"],
        ["厨电", "方太、老板", "美的、华帝", "3000-8000元", "风量风压、清洁便利性"],
    ]
)
add_source("数据来源：2025-2026年行业监测数据及多渠道达人推荐综合整理")

add_h3("热门型号参考")

add_para("电视：")
add_para("海信E5Q Pro（75寸约4500-5000元）：墨晶屏抗反光，信芯AI画质芯片")
add_para("TCL T7L Ultra（75寸约4500元，85寸约6000元）：蝶翼星耀屏，万元内画质标杆")
add_para("小米Redmi X2026（85寸约4600元）：Mini LED，2400尼特峰值亮度")

add_para("冰箱：")
add_para("容声526：双系统双循环，一级净味，一键自动制冰")
add_para("美的550：法式五门，双系统双蒸发器")
add_para("卡萨帝550：MRA低氧窖藏，智慧动态杀菌")

add_para("洗衣机：")
add_para("小天鹅小乌梅3.0（约3000元）：超微气泡冷水洗，银离子除菌")
add_para("海尔云溪4.0：FPA直驱电机，精华洗")
add_para("卡萨帝朗境W5（万元以上）：双动力热泵系统")
add_source("数据来源：京东热销数据及多渠道达人推荐验证")

add_h3("国补叠加攻略")

add_image("T5_国补攻略.png")

add_para("2026年国家家电补贴政策仍在执行：6类家电1级能效按售价15%补贴，每件最高1500元。更划算的做法是优先买8000-10000元档位的产品，补贴力度最大。一台8000元的Mini LED电视，国补省1200元，实际到手6800元。空调1.5匹线上均价已降到2100元，补贴再省300-600元。")
add_source("数据来源：2026年国家家电以旧换新补贴政策")

add_h3("三星清仓能不能捡漏？")

add_para("不建议。Tizen系统生态远不如国产安卓TV，Mini LED技术已被国产拉开差距，白电本身在市占率不到1%的时候产品力就已经落后了，停售后零部件长期供应也存在不确定性。同价位国产品牌配置更优、售后更稳，没必要冒险。")

# ── Closing ──
add_h2("写在最后")

add_para("过去十二年，11个外资家电品牌先后退出或大幅收缩中国市场。这不是偶然，是中国制造业从“代工组装”到“技术输出”的必然结果。")

add_para("从京东方的面板到海信的画质芯片，从小米的AIoT生态到卡萨帝的高端范式，中国家电品牌用了二十年走完了从追赶到超越的全过程。外资品牌撤退之后留下的不是空白，而是一个更透明、竞争更充分、对消费者更友好的市场。")

add_para("当然，也有少数外资品牌在特定领域仍有优势——大金的中央空调、戴森的吹风机、博西的嵌入式洗碗机。但在中国家庭最核心的家电品类上，国产品牌已经不是“凑合能用”，而是在同价位提供了更好的技术和体验。")

add_para("买什么，市场已经给出了答案。")

# ── Save ──
out_path = os.path.join(OUTPUTS, "外资家电集体撤退_终稿.docx")
doc.save(out_path)
print(f"Word文档已生成: {out_path}")
print(f"文件大小: {os.path.getsize(out_path) / 1024 / 1024:.1f} MB")
