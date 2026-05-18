#!/usr/bin/env python3
"""生成 三星家电退出中国 头条深度图文 Word 文档 · 深度去AI味版"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUTS = os.path.join(BASE_DIR, "outputs")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style_normal = doc.styles['Normal']
style_normal.font.name = 'Microsoft YaHei'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.line_spacing = 1.8
style_normal.paragraph_format.space_after = Pt(6)
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


def add_source(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_image(filename, width_inches=6.0):
    path = os.path.join(OUTPUTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
    else:
        add_para(f"[图片缺失: {filename}]")


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Microsoft YaHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Microsoft YaHei'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.add_paragraph()


# ════════════════════════════════════════════
#  ARTICLE CONTENT  (去AI味终稿)
# ════════════════════════════════════════════

add_title("三星家电被中国赶走了：34年来它做错了什么？")

add_image("T5_封面图.png", 5.2)

add_para("5月6日，三星电子在官网挂出一条不到200字的公告：即日起，在中国大陆市场停止销售所有家电产品。电视、冰箱、洗衣机、空调、吸尘器、空气净化器——13个品类，一刀切干净。手机不受影响，继续卖。")

add_para("一家从1992年入华、深耕了34年的跨国巨头，最后只留下一句“售后服务照常”。")

add_para("但真正让人意外的不是退出本身，而是同一天的另一组数字：三星股价暴涨14.41%，市值冲破1万亿美元，成为继台积电之后第二家跻身“万亿俱乐部”的亚洲公司。KOSPI指数历史性地突破7000点，韩国股市进入“7000点时代”。")

add_para("退出全球最大的家电消费市场，股价不跌反涨。听起来很反直觉，但这背后是一段中国家电产业花了34年才写完的逆袭故事。")

# ── Section 1 ──
add_h2("三组数据：三星从“霸主”到“不足1%”")

add_para("先把时间拉回2013年。那一年，三星电视在中国市场的占有率接近18%，稳居榜首。彩电一年卖出255万台，销售额约30亿美元。三星的标志出现在每一个高端百货商场的显眼位置，买三星就是“体面”的代名词。")

add_para("然后是断崖。")

add_para("奥维云网的数据显示，截至2026年4月，三星家电在中国线下渠道的市占率已经惨不忍睹：彩电3.62%，排第5；冰箱0.41%，排第14；洗衣机0.38%，排第15。冰箱和洗衣机的份额加在一起，都不到1%。")

add_table(
    ["品类", "2013年巅峰市占率", "2026年4月市占率", "当前排名"],
    [
        ["彩电", "约18%", "3.62%", "第5"],
        ["冰箱", "行业前列", "0.41%", "第14"],
        ["洗衣机", "行业前列", "0.38%", "第15"],
    ]
)
add_source("数据来源：奥维云网（AVC）监测数据")

add_image("T5_三星市占率断崖.png")

add_para("第二组数据更能说明问题。打败三星的不是市场萎缩，而是国产品牌的集体反超。2025年，中国电视市场前八大品牌——海信、TCL、小米、创维、长虹、海尔、康佳、华为——全年出货量达3096.3万台，合计份额94.1%。外资品牌（三星、索尼、飞利浦、夏普）加在一起，出货量不到100万台，份额不到5%。")
add_source("数据来源：洛图科技（RUNTO）《中国电视市场品牌出货月度追踪》")

add_para("第三组数据解释了三星为什么“不挣扎就走”。2026年第一季度，三星合并营收133.9万亿韩元，营业利润57.2万亿韩元，同比暴增756%。但拆开看结构，半导体（DS）部门一家贡献了53.7万亿韩元利润，占全公司93.4%。而负责电视和生活家电的VD+DA部门，2025年全年亏损2000亿韩元（约9.26亿元人民币）——这是三星家电入华34年来头一次亏损。")
add_source("数据来源：三星电子2026年Q1财报")

add_para("说白了，三星不是被打趴的，是算完账自己走的。一边是卖一颗HBM4芯片赚70%利润率、供不应求；另一边是在中国卖冰箱，扣掉渠道费和售后成本，一台赚几百块。2026年三星规划了110万亿韩元（约733亿美元）资本开支，全砸向半导体。家电这条线，不陪了。")

add_image("T5_半导体芯片.jpg", 5.5)

# ── Section 2 ──
add_h2("34年三个致命失误")

add_para("三星的败退不是一夜之间的事。把时间线拉长看，三个战略失误一步步把三星推向了今天这个结局。")

add_h3("失误一：本土化严重缺失，决策权死死攥在首尔。")

add_para("中国电子视像行业协会秘书长董敏说得直白：“三星家电在中国市场的定价权和管理决策权都在韩国总部。”过去十年，中国家电渠道从苏宁国美大卖场，跳到京东天猫，再跳到抖音直播带货。国产品牌能按周甚至按天调整策略和价格，三星的跨国决策链走一个审批流程就错过了整个促销季。")

add_para("产品层面更明显。中国家庭需要超大容量多开门冰箱、适合小户型公寓的超薄全嵌洗衣机，三星还在推全球统一的标准款。同等甚至落后的配置，价格却比国产高出一大截。品牌溢价在性价比面前撑不住了。")
add_source("数据来源：澎湃新闻、中国经营报报道")

add_h3("失误二：供应链脱钩，从“面板霸主”沦为“组装厂”。")

add_para("三星电视能称霸全球十几年，靠的是从LCD面板到整机的垂直一体化优势。但近几年，京东方、华星光电等中国面板厂在LCD领域产能和成本都把三星甩在了后面，三星选择战略性退出LCD面板制造，转向OLED。")

add_para("这个决定在集团层面说得通，但对电视终端业务是致命的。退出面板制造后，三星在主流彩电上不得不向京东方、华星光电等中国厂商采购核心零部件。当昔日的全球霸主沦为要向竞争对手买屏幕的“组装厂”，成本优势彻底没了。打价格战的时候，三星电视完全扛不住拥有供应链腹地优势的国产玩家。")
add_source("数据来源：T客邦、21经济报道")

add_h3("失误三：错失AIoT生态，SmartThings成了“信息孤岛”。")

add_para("现在82%的中国消费者买家电会优先看智能联动功能，91%期待远程控制和语音交互。小米有米家和小爱同学，华为有鸿蒙全屋智能，海尔有三翼鸟——国产品牌早就把家电绑进了互联网生态。")

add_para("三星不是没有SmartThings平台，但在中国市场，受限于本土互联网服务的接入能力、数据合规要求和生态合作伙伴的匮乏，三星的智能家电成了无法互联互通的孤岛。消费者习惯了用一句话控制全屋灯光、窗帘、空调和电视，谁还愿意用三个遥控器分别操作一台“聪明但不合群”的三星电视？")
add_source("数据来源：GfK中怡康2026年调研")

add_image("T5_智能家居厨房.jpg", 5.5)

# ── Section 3 ──
add_h2("国产三条路线：技术突破、价格颠覆、高端夺权")

add_para("三星输了，但赢它的不是某一个对手，而是一整个生态体系的超越。三条路线，每一条都扎在三星的软肋上。")

add_image("T5_三条国产路线.png")

add_h3("路线一：海信和TCL，在三星的主场上打赢了它。")

add_image("T5_现代客厅电视.jpg", 5.5)

add_para("TCL背靠华星光电的面板垂直整合能力，海信死磕自研信芯画质处理芯片，两家共同在Mini LED赛道上实现了对三星的弯道超车。2024年，TCL在Mini LED出货量上反超三星；2025年，三星在Mini LED出货量仅排全球第四。")

add_para("更具标志性的转折发生在2025年12月——TCL单月电视出货量首次超越三星，登顶全球榜首。2025年全年，三星3530万台（连续20年全球出货量居首），TCL 3040万台，差距缩小到490万台。海信紧随其后，2990万台。")

add_para("价格对比更残酷：同规格Mini LED产品，国产品牌普遍便宜三到五成。海信Mini LED入门款4599元，三星同技术产品最低也要11000元以上。在0到4000元这个中国家庭最集中的价格区间，三星几乎没有产品覆盖。")

add_table(
    ["2025年全球电视出货TOP3", "出货量", "核心优势"],
    [
        ["三星", "3530万台", "全球品牌惯性，QLED"],
        ["TCL", "3040万台", "华星光电+Mini LED+百寸巨幕"],
        ["海信", "2990万台", "自研画质芯片+激光电视+多品牌矩阵"],
    ]
)
add_source("数据来源：奥维睿沃（AVC Revo）2025年全球彩电市场报告")

add_image("T5_全球电视出货对比.png")

add_h3("路线二：小米，用生态把家电变成了“入口”。")

add_para("小米在家电领域的逻辑跟传统厂商完全不同——它不是在卖硬件，是在铺设AIoT生态的触点。2025年，小米空调出货超850万台（同比增24%），冰箱超280万台（同比增4%），洗衣机超230万台（同比增18%），三大白电均创历史新高。")

add_para("背后的护城河是生态数据：AIoT平台连接设备10.79亿台，五件以上设备用户2270万，米家App月活1.127亿。当一个用户家里已经有小米电视、小爱音箱、扫地机器人，他下一台空调大概率还是小米——不是因为参数多强，而是因为“一句话就能联动”。这种生态粘性，传统外资品牌复制不了。")
add_source("数据来源：小米集团2025年财报")

add_h3("路线三：卡萨帝，在三星最骄傲的高端市场完成了夺权。")

add_para("过去高端家电是三星、西门子、松下的“自留地”。海尔旗下的卡萨帝把这条鄙视链打碎了。2026年1-14周GfK中怡康数据显示，卡萨帝在1.5万以上的高端市场份额：冰箱71.1%，洗衣机约80%，空调44.7%。洗烘一体2万+价位段，份额突破95%。")

add_para("以最高均价拿最高份额——不是靠便宜赢的，是靠产品和体验在同一个价位段正面击败了外资品牌。卡萨帝指挥家套系还拿到了中家院L4级智慧家电认证，是高端家电领域头一个获此殊荣的品牌。市面上大多数智能家电还停留在L2、L3级的“被动响应”，卡萨帝已经做到了“能听、会看、会思考”的主动服务。")
add_source("数据来源：GfK中怡康2026年1-14周监测数据、海尔智家2026年一季报")

# ── Section 4 ──
add_h2("走了一个三星，来了一个更大的时代")

add_para("三星的退出不是孤立事件，而是外资家电在中国市场集体退潮的一波高峰。三洋2014年被惠而浦收购、东芝2018-2021年逐步剥离、飞利浦2020年以44亿美元把家电业务卖给高瓴资本、夏普2016年被富士康收购后持续边缘化、LG 2024年底全线撤出中国大陆、索尼2025年把电视业务托管给TCL——一份名单，就是过去三十年中国家电竞争史的失败者花名册。")
add_source("数据来源：各品牌公开公告及行业研报")

add_image("T5_外资撤退时间线.png")

add_para("有意思的是，并不是所有外资品牌都在跑。松下在2026年反而在加码中国——完成集团重组后，决策权下放到“中国东北亚事业”板块，口号是“中国的事情在中国决定”。削减低毛利产品，聚焦“住空间解决方案”，直营品类销售同比翻倍。")

add_para("这说明一件事：愿意放下跨国企业的历史包袱、深度融入中国本土供应链和数字化生态的，在中国庞大且分层的高端市场里，依然有活路。三星选择走，是因为它有一个利润暴增93%的半导体业务作为退路——其他没有这条退路的品牌，只能学着适应。")

# ── Section 5 ──
add_h2("消费者最关心的两件事")

add_para("三星退了，很多人最先想到的问题是：已经买的三星家电怎么办？以后买家电选什么？")

add_para("售后方面，三星在公告中明确承诺，将继续按照《消费者权益保护法》和国家三包规定提供售后服务，官方认证的售后服务中心、400热线和“三星服务”微信小程序保持正常运营。京东也跟进宣布，为过保产品提供免费上门检测，维修价值不高时推荐以旧换新。简单说，已经买的不用慌。")

add_para("至于“买什么”，2026年的答案是：放下洋品牌滤镜，按品类选国产。")

add_table(
    ["品类", "首选", "备选", "关键技术点"],
    [
        ["电视", "海信、TCL", "小米", "Mini LED背光分区数、HDR峰值亮度"],
        ["冰箱", "海尔/卡萨帝", "容声、美的", "双系统双循环、净味技术"],
        ["洗衣机", "海尔/卡萨帝", "小天鹅、美的", "直驱电机、精华洗"],
        ["空调", "美的", "格力、海尔", "APF能效值、实际体验"],
    ]
)
add_source("数据来源：2025-2026年行业监测数据综合整理")

add_image("T5_国产替代选购指南.png")

add_para("还有几个细分领域，外资品牌仍有优势：大金在中央空调和VRV系统上仍然强势，戴森的吹风机和吸尘器目前没有真正的国产平替，博西家电在高端嵌入式洗碗机和洗烘一体机上仍有竞争力。")

# ── Closing ──
add_h2("写在最后")

add_para("三星家电在中国的34年，是一个外资品牌从“仰视对象”变成“被淘汰者”的完整样本。它不是被某一家企业打败的，是被一整个产业链的升级打败的——从京东方的面板到海信的画质芯片，从小米的AIoT生态到卡萨帝的高端范式，中国家电用了二十年完成了从“组装代工”到“技术输出”的转身。")

add_para("但也要看到另一面。三星退出中国家电市场的那天，市值站上了1万亿美元。半导体业务一个季度赚的钱，比华为干两年半还多。这不是一个“输了”的故事，这是一个“算清楚了，不跟你耗了”的故事。")

add_para("对中国消费者来说，三星退场带来的不是恐慌，是一个更清晰的现实：在绝大多数家电品类上，国产品牌已经不是“凑合能用”，而是在同价位提供了更好的技术和体验。选什么，市场已经给出了答案。")

# ── Save ──
out_path = os.path.join(OUTPUTS, "三星家电被中国赶走了_终稿.docx")
doc.save(out_path)
print(f"Word文档已生成: {out_path}")
print(f"文件大小: {os.path.getsize(out_path) / 1024 / 1024:.1f} MB")
