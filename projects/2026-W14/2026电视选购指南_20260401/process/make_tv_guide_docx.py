from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置中文字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading_custom(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def add_para(doc, text):
    return doc.add_paragraph(text)

def add_pic_center(doc, pic_path, caption="", width_inches=6):
    doc.add_picture(pic_path, width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)

# ========== 封面 ==========
doc.add_picture('outputs/TV_封面图.png', width=Inches(6.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ========== 一、先定人群再定机型 ==========
add_heading_custom(doc, '一、先定人群，再定机型：你的需求属于哪类？', 1)

doc.add_paragraph('很多人买电视的第一步就错了。不是先看预算，而是先问"哪款最好"。')

doc.add_paragraph('好电视确实多，但"最好"的只有一台：最适合自己的那台。')

doc.add_paragraph('在掏钱之前，先问自己三个问题：')

p = doc.add_paragraph()
p.add_run('第一，你的核心使用场景是什么？').bold = True
doc.add_paragraph('是晚上关灯看电影，还是白天一家人围坐刷剧？是接PS5打主机，还是就看看新闻联播？')

p = doc.add_paragraph()
p.add_run('第二，你的客厅光线条件如何？').bold = True
doc.add_paragraph('朝南大落地窗的客厅和能拉上窗帘的独立影音室，选电视的技术路线完全相反。前者需要抗光能力强、亮度高的Mini LED，后者更适合像素级控黑的OLED。')

p = doc.add_paragraph()
p.add_run('第三，你的智能生态是哪家的？').bold = True
doc.add_paragraph('米家用户、华为全家桶用户和没有生态负担的用户，各自的最优解完全不同。')

doc.add_paragraph('想清楚这三个问题，你已经排除了80%的错误选项。')

doc.add_paragraph('我把人群分成四类，对号入座：')

add_pic_center(doc, 'outputs/TV_选购决策图.png', '图1：2026电视选购人群分类与预算参考', 6.5)

# ========== 二、分预算推荐 ==========
add_heading_custom(doc, '二、分预算推荐：三个价位段精选机型', 1)

add_heading_custom(doc, '5000元以下：入门之选，够用就好', 2)
doc.add_paragraph('这个价位别追求什么极致画质，但"够用"二字要做到位并不简单。2026年入门市场的坑比高端市场还多。')

p = doc.add_paragraph()
p.add_run('核心原则：').bold = True
p.add_run('RGB真4K（绝对不碰RGBW伪4K）、峰值亮度≥800nit、原生60Hz或以上、有MEMC运动补偿、内存4GB+64GB起步。')

add_pic_center(doc, 'outputs/02_modern_tv.jpg', '图2：入门电视也能有好画质', 5)

add_heading_custom(doc, '5000-20000元：主流战场，卷王扎堆', 2)
doc.add_paragraph('这是80%家庭的选择区间，也是2026年竞争最激烈的价位段。Mini LED技术在这个价位已经全面普及。')

add_pic_center(doc, 'outputs/TV_技术对比图.png', '图3：MiniLED vs OLED vs QLED vs LED 技术对比', 6.5)

add_heading_custom(doc, '20000元以上：高端玩家的殿堂', 2)
doc.add_paragraph('这个价位段的用户已经不是"买电视"了，而是在"买信仰"或者"买极致体验"。选的不是参数，是调校哲学。')

add_pic_center(doc, 'outputs/01_tv_living_room.jpg', '图4：高端电视带来影院级体验', 5)

# ========== 三、分用途推荐 ==========
add_heading_custom(doc, '三、分用途推荐：你的场景，你的最优解', 1)

add_heading_custom(doc, '影音发烧友', 2)
add_para(doc, '核心需求：黑位表现、色彩精准度、暗场细节、HDR高光还原。')
add_pic_center(doc, 'outputs/04_tv_tech.jpg', '图5：影音发烧友追求极致画质', 5)

add_heading_custom(doc, '游戏玩家', 2)
add_para(doc, '核心需求：刷新率、输入延迟、VRR防撕裂、HDR游戏表现。')
add_pic_center(doc, 'outputs/05_gaming_setup.jpg', '图6：游戏玩家需要高刷低延迟', 5)

add_heading_custom(doc, '智能家居用户', 2)
add_para(doc, '核心需求：生态联动流畅、投屏体验、语音助手。')
add_pic_center(doc, 'outputs/03_remote_control.jpg', '图7：智能电视成为家庭控制中枢', 5)

add_heading_custom(doc, '追求家居美学', 2)
add_para(doc, '核心需求：超薄设计、贴墙安装、艺术模式。')

# ========== 四、技术扫盲 ==========
add_heading_custom(doc, '四、技术扫盲：OLED/Mini LED/传统LED，到底怎么选？', 1)

add_heading_custom(doc, 'Mini LED：2026年绝对主流', 2)
add_para(doc, 'Mini LED本质上是LCD的"超级进化版"。用数千颗微缩LED灯珠做背光源，配合精密的分区控光，大幅缩小了与OLED的画质差距。')

add_para(doc, '2026年Mini LED的渗透率已达45.8%，其中85英寸以上机型占比超70%。买55英寸以上客厅电视，Mini LED是当前最优选。')

add_pic_center(doc, 'outputs/TV_品牌矩阵图.png', '图8：2026年主流电视品牌技术阵营', 6.5)

add_heading_custom(doc, 'OLED：黑位无敌，但需理性看待', 2)
add_para(doc, 'OLED最大的优势是像素级自发光。每个像素自己发光，自己控制亮度。这意味着在显示黑色时，像素完全熄灭，是真正的"纯黑"。')

add_heading_custom(doc, 'Micro LED：看看就好', 2)
add_para(doc, '这是2026年技术的天花板，但价格嘛，TCL Max163M售价249999元。顶级豪宅专属，咱们普通人就当看个热闹。')

# ========== 五、闭坑指南 ==========
add_heading_custom(doc, '五、闭坑指南：商家不会告诉你的5个真相', 1)

add_heading_custom(doc, '1. 灯珠数≠分区数', 2)
add_para(doc, '这是2026年最常见的营销陷阱。部分机型标注"10000颗灯珠"，但实际物理控光分区只有200个。选购Mini LED，只认物理分区数。')

add_heading_custom(doc, '2. "瞬时峰值亮度"是智商税', 2)
add_para(doc, '有些机型标注"1200nit峰值亮度"，但实测持续亮度只有400-500nit，开机一分钟后直接跳水。')

add_heading_custom(doc, '3. DLG插帧伪高刷会废掉一半清晰度', 2)
add_para(doc, 'DLG技术是通过芯片"猜帧"生成中间画面，看起来刷新率提高了，但实际会损失50%原生像素。')

add_heading_custom(doc, '4. HDMI 2.1有"残血"和"满血"之分', 2)
add_para(doc, '游戏玩家必须选支持FRL通道的满血HDMI 2.1接口。')

add_heading_custom(doc, '5. 2+32GB的电视用两年就卡', 2)
add_para(doc, '无论预算多低，4GB+64GB是起步线。')

add_pic_center(doc, 'outputs/TV_闭坑指南图.png', '图9：2026电视选购5大陷阱', 6.5)

# ========== 六、购买时机 ==========
add_heading_custom(doc, '六、购买时机：什么时候买最划算？', 1)

add_para(doc, '四大黄金节点：618、双11、年货节（春节前）、818。')

add_para(doc, '其中618和双11是力度最大的两个节点，叠加15%国补（单件最高1500元），很多机型可以省下上千元。')

add_para(doc, '国补政策现在是常态。单件最高1500元15%补贴，没有截止时间限制。如果你正好有换机需求，现在就是最好的时机。')

# ========== 七、一张表 ==========
add_heading_custom(doc, '七、一张表告诉你：不同人群的最优解', 1)

add_pic_center(doc, 'outputs/TV_选购决策图.png', '图10：2026电视选购终极推荐表', 6.5)

# ========== 写在最后 ==========
add_heading_custom(doc, '写在最后', 1)

doc.add_paragraph('2026年的电视市场，选择比以往任何时候都多，坑也比以往任何时候都深。')

doc.add_paragraph('但核心逻辑从来没有变过：先想清楚自己的需求和场景，再去匹配机型，而不是被参数牵着鼻子走。')

doc.add_paragraph('一台电视少则用三五年，多则七八年。与其追逐最新的技术名词，不如多花半小时想清楚自己真正需要什么。')

p = doc.add_paragraph()
p.add_run('记住这个顺序：').bold = True
p.add_run('人群定位 → 场景需求 → 预算范围 → 技术选型 → 具体型号。')

doc.add_paragraph('按这个顺序走，你大概率不会选错。')

p = doc.add_paragraph()
p.add_run('最好的电视，不是参数最高的那一台，而是最适合你生活场景的那一台。').bold = True

doc.add_paragraph('祝你选到心仪的那台电视。')

# 数据来源
doc.add_paragraph()
p = doc.add_paragraph('数据来源：本文参考了2026年主流品牌官方参数、第三方测评机构实测数据（RTINGS、中关村在线、太平洋电脑网等）及中国电子视像行业协会技术报告。文中价格均为电商平台日常活动价叠加15%国补后的参考价，实际价格以购买时平台显示为准。')

doc.save('outputs/2026电视选购全景指南_配图版.docx')
print('Done: 2026电视选购全景指南_配图版.docx')
