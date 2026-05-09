from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置中文字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_bold_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

# 标题
title = doc.add_heading('Rtings 8.4分的TCL旗舰，国内6799元就能抱回家？', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Rtings背景
add_bold_para(doc, 'Rtings 8.4分，这个数字什么水平？')
doc.add_paragraph('先说背景。Rtings.com，北美最严苛的电视评测机构，测试标准化到有点变态，评分出了名的保守。能上8分的机器，基本就是全球公认的第一梯队了。')
doc.add_paragraph('最新一期榜单里，TCL QM8K拿了8.4分，和索尼Bravia 9、海信U8QG站一排。更离谱的是，它的综合画质评分还把LG G4和三星S95D这两款前代旗舰OLED给压下去了——这俩在北美卖得比QM8K还贵。')
doc.add_paragraph('Mini LED阵营头一回在硬核评测场子上，正面撕开了OLED的防线。这是QM8K的底气，也是我们后面聊Q9M Pro的起点。')
doc.add_paragraph('还有一件事更让人意外：这台在北美卖7200块的机器，国内6799块就能拿下。而且国内版还多了几项北美版没有的功能。')

# 第一张配图
doc.add_picture('outputs/01_中美价格对比.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('图1：中美TCL旗舰电视价格对比（含65寸/85寸）')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 价格部分
doc.add_heading('先把价格说清楚', level=2)
doc.add_paragraph('科技圈有句话流传挺广："北美TCL旗舰卖一万六，国内只要五千多。"')
doc.add_paragraph('说对了一半。85寸确实，北美QM8K约15800块，国内Q9M Pro约11999块，差了3900块，这没毛病。')
doc.add_paragraph('但65寸才是大多数人在看的尺寸——')
p = doc.add_paragraph()
p.add_run('北美QM8K：').bold = True
p.add_run('$999.99，大概7200块。')
p = doc.add_paragraph()
p.add_run('国内Q9M Pro：').bold = True
p.add_run('国补后6799块。')
doc.add_paragraph('差了不到400块，不到6%。说"便宜了一半"？差太远了。说白了，就是花差不多的钱，买到了更好的东西。')

# Rtings评分维度
doc.add_heading('Rtings 8.4分怎么来的？5个维度拆解', level=2)

items = [
    ('对比度', 'Mini LED千级分区加上VA面板的原生高对比度。QM8K实测原生对比度超过8000:1。暗室里，纯黑画面点亮一颗星，只有那一像素是亮的，周围全黑——能做到这个效果，和OLED一个水准。'),
    ('HDR峰值亮度', '5000尼特。HDR内容里太阳光、火焰、金属反光这些高频高亮细节，能几乎不压缩地还原出来。OLED那边亮度有天花板，画面一亮就得拼命压亮度，细节丢得多。QM8K这边不需要太激进，高光细节保留得多很多。'),
    ('色彩容积', 'TCL的SQD技术厉害在哪？亮度拉满的时候，色彩饱和度不会跟着掉。看HDR大片里那些饱和度很高的颜色，不会感觉"一亮起来颜色就变浅了"。'),
    ('EOTF曲线追踪', 'QM8K在处理暗部信号的时候，曲线追踪比堆料更猛的QM9K还顺。没有QM9K那种轻微的"暗部死黑"，暗部的灰阶过渡更自然，更接近导演想让你看到的效果。这是RTINGS觉得QM8K比QM9K更值得买的关键原因。'),
    ('游戏性能', '144Hz原生刷新率、VRR、ALLM，输入延迟13毫秒以内，Game Accelerator 288还能跑到288Hz。你拿PS5或者Xbox Series X跑《黑神话》，画面的跟手程度和高速移动时的清晰度，60Hz电视根本给不了。'),
]
for name, desc in items:
    p = doc.add_paragraph()
    p.add_run(f'{name}：').bold = True
    p.add_run(desc)
doc.add_paragraph('五个维度没有明显短板。8.4分就是这么来的。')

# 第二张配图
doc.add_picture('outputs/02_SQD技术原理.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('图2：SQD-Mini LED技术原理与RGB-Mini LED对比')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# SQD技术
doc.add_heading('SQD技术：TCL真正厉害的地方在这里', level=2)
doc.add_paragraph('QM8K和Q9M Pro能拿高分，除了分区数和亮度，更重要的是色彩稳定性和精准度。而这两点，都要靠SQD-Mini LED技术。')
doc.add_paragraph('RGB-Mini LED有个无解的问题：红、绿、蓝三种颜色的LED芯片在高功率发光时，热衰减曲线完全不一样——红光衰减最快。你用个一年半载，屏幕就会出现不可逆的色偏，鲜红变橙红，翠绿变黄绿。更要命的是，亮度一拉满，密集排列的RGB三色光源互相干扰，高亮场景下颜色饱和度直接崩掉。')
doc.add_paragraph('TCL的SQD方案换了个思路：不用RGB三色混光，直接用一颗纯蓝光白光LED作为发光基底，上面加一层纳米级超量子点滤光膜。蓝光穿过这层膜，激发出来的红光和绿光谱带极窄，纯度极高。物理路径短，损耗小，颜色准。')

p = doc.add_paragraph()
p.add_run('落到你看片的时候，感受差在哪？').bold = True
doc.add_paragraph('看《沙丘2》厄拉科斯的日落，普通电视高光刺眼、暗部死黑，HDR效果被粗暴的色调映射破坏得七七八八。Q9M Pro上，沙漠里阳光的温暖金色和阴影中香料粉末的微弱橙色同时存在、层次分明，高光不过曝，暗部不死黑。')
doc.add_paragraph('看《蓝色星球》深海，普通电视黑乎乎一片，什么细节都看不清。Q9M Pro上，生物的微弱磷光被精准点亮，周围干净利落，没有灰蒙蒙的背光散射。')
doc.add_paragraph('玩HDR赛车游戏，过隧道的时候内外明暗剧烈切换，普通LCD会有明显的光晕拖尾。SQD控制逻辑简单高效，这种现象大幅减轻。')

p = doc.add_paragraph()
p.add_run('100% BT.2020色域覆盖').bold = True
p.add_run('。这是SQD技术的终极证明。现在市面上顶级的OLED和RGB-Mini LED，在BT.2020面前覆盖率基本卡在90%到97%之间，上不去。Q9M Pro和T7M Pro借助SQD达到了100%。')
doc.add_paragraph('以前这种能力只有专业级监视器才有。现在T7M Pro告诉你，6000块级，就能有。')

# 第三张配图
doc.add_picture('outputs/03_Q9M_QM8K雷达对比.png', width=Inches(6))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('图3：Q9M Pro与QM8K核心参数雷达对比')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 硬数据
doc.add_heading('硬数据：Q9M Pro不是阉割版', level=2)
doc.add_paragraph('3552对3800，5000尼特对5000尼特。28个分区的差别就是屏幕尺寸不同导致的背光板面积差，正常工程公差，不影响你看片的感受。')
p = doc.add_paragraph()
p.add_run('国内版多出来的东西才是值回差价的部分：').bold = True
doc.add_paragraph('300Hz超频刷新率——北美版没有。你拿RTX 5090跑《CS2》或者《永劫无间》，300Hz的跟手感，144Hz真的给不了。')
doc.add_paragraph('TSR AI画质引擎——专门给国内视频生态优化的。腾讯视频、爱奇艺为了省带宽，把4K压成什么德行大家心里有数。TSR引擎在本地实时做像素级修复，在线视频的观感能拉到接近蓝光4K。')
doc.add_paragraph('灵控系统3.0——没有开机广告。点一下遥控器，两秒进桌面。算下来每天省30秒，一年轻松多出三个小时。')
doc.add_paragraph('安桥Onkyo 2.1.2 Hi-Fi——北美版是B&O，走轻奢路线；国内换成安桥，声底厚实、人声解析力强、低频下潜有力，还有杜比全景声天空声道。对中国家庭影院玩家的口味，这套方案明显拿捏得更准。')
doc.add_paragraph('65寸差价不到400块，功能反而更多。85寸呢？差价5800块。')

# 65还是85
doc.add_heading('65寸还是85寸？直接给结论', level=2)
p = doc.add_paragraph()
p.add_run('65寸：买Q9M Pro不是因为便宜，是因为更值。').bold = True
doc.add_paragraph('差价不到400块，但换来了无广告系统、TSR画质引擎、300Hz超频刷新率。说是功能升级没问题，说是阉割补偿就太委屈它了。')
doc.add_paragraph('T7M Pro 65寸约5270块起，1152个分区、2200尼特亮度，同样100% BT.2020色域和安桥音响。你不是极致参数党的话，这机器就是6000块以内最理性的选择。')
p = doc.add_paragraph()
p.add_run('85寸：差价才是真正的大头。').bold = True
doc.add_paragraph('85寸差价5800块。这钱你可以拿来买一套高端Soundbar，加一套旗舰挂架，再加三年有线电视费，还有剩。')
doc.add_paragraph('国内版85寸Q9M Pro同样3552个分区、5000尼特、SQD技术、安桥音响。这5800块差价，买到的是和北美版几乎一模一样的硬件。')

# T7M Pro
doc.add_heading('T7M Pro：把高端门槛直接拉低了一半', level=2)
doc.add_paragraph('最后单独说说T7M Pro，因为这机器的意义不只是另一个选项。')
doc.add_paragraph('1152个分区、2200尼特亮度、100% BT.2020色域、60mm超薄机身、安桥音响、灵控系统3.0无广告——65寸国补后约5270块。')
doc.add_paragraph('以前，1000个分区加2000尼特亮度，是"真HDR体验"的及格线，门槛基本在万块以上。')
doc.add_paragraph('T7M Pro把这条线拉到了5270块。')
doc.add_paragraph('这不是打价格战，是TCL靠华星光电面板和自己的垂直供应链整合，在重新定义高端电视该卖什么价。')

# 结论
doc.add_heading('结论', level=2)
doc.add_paragraph('北美花7200块买QM8K，国内花6799块买Q9M Pro，体验差多少？')
p = doc.add_paragraph()
p.add_run('说实话，国内版好几个地方还更好。').bold = True
doc.add_paragraph('一样的3552分区、5000尼特光控底座，额外加上300Hz超频、TSR引擎、无广告系统、安桥音响——6799块的Q9M Pro，体验比7200块的北美QM8K只好不差。')
doc.add_paragraph('"便宜了一半"这话不成立。但"花更少的钱，体验反而更好"，是真的。')
doc.add_paragraph('85寸那5800块的差价，是这句话最扎实的证据。')
doc.add_paragraph('RTINGS的编辑们还在琢磨QM8K和索尼Bravia 9哪个更值的时候，咱们这边只需要想一件事：Q9M Pro，还是T7M Pro？')
p = doc.add_paragraph()
p.add_run('这道选择题，本身就挺幸福的。').bold = True

# 数据来源
doc.add_paragraph()
p = doc.add_paragraph('数据来源：Rtings.com综合评分榜及评测报告、TCL官方技术白皮书、新浪财经/Gizmochina/NotebookCheck/HDTVTest等科技媒体评测报道')

doc.save('outputs/TCL_Q9M_Pro_T7M_Pro深度横评_配图版.docx')
print('Done')
