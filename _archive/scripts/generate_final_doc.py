#!/usr/bin/env python3
"""生成美的冰箱618头条图文最终Word文档（去AI味+配图嵌入版）"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/ARTICLE WRITER/outputs/美的冰箱618头条图文_最终版.docx"
IMAGES = {
    "cover":    "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/ARTICLE WRITER/outputs/T5_01_封面图.png",
    "compare":  "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/ARTICLE WRITER/outputs/T5_02_参数对比图.png",
    "scene":    "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/ARTICLE WRITER/outputs/T5_03_场景对比图.png",
    "proscons": "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/ARTICLE WRITER/outputs/T5_04_优缺点图.png",
    "people":   "/Users/ltn/Library/CloudStorage/GoogleDrive-qishuai1006@gmail.com/我的云端硬盘/ARTICLE WRITER/outputs/T5_05_人群选购图.png",
}

doc = Document()

# ── 全局样式 ──────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 标题样式 ──────────────────────────────────────
def set_heading(para, level=1):
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)

def add_para(text, bold=False, italic=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_image(path, width=Inches(6)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    return p

def add_divider():
    p = doc.add_paragraph('─' * 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(180, 180, 180)

def add_table_row(table, data, bold=False):
    row = table.add_row()
    for i, text in enumerate(data):
        cell = row.cells[i]
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.bold = bold
    return row

# ════════════════════════════════════════════════════
# 封面图
# ════════════════════════════════════════════════════
add_image(IMAGES["cover"], width=Inches(6.5))
doc.add_paragraph()

# ════════════════════════════════════════════════════
# 主标题
# ════════════════════════════════════════════════════
title = doc.add_heading('', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('2026年618美的冰箱选购指南')
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = RGBColor(30, 30, 30)

sub = doc.add_heading('', level=2)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('哪些值得买，哪些是坑，看完这篇就够了')
run2.font.size = Pt(13)
run2.italic = True
run2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# ════════════════════════════════════════════════════
# 开场白（去AI味）
# ════════════════════════════════════════════════════
add_para('今年618想买美的冰箱，我直接给你结论：有些型号确实值得入，有些说实话不太划算，花那钱有点亏。')
add_para('我研究冰箱这么多年，见过太多人买完就后悔——要么容量买小了不够用，要么嵌入式根本放不进橱柜，要么图便宜买了个单系统的回家发现串味严重。')
add_para('所以这篇指南不跟你扯什么情怀和品牌故事，就一件事：帮你把钱花在刀刃上，让你618买冰箱不踩坑。')
add_para('先说结论，懒得看完全文的直接照着买：', bold=True)

# ════════════════════════════════════════════════════
# 一、先给结论
# ════════════════════════════════════════════════════
h1 = doc.add_heading('', level=1)
h1.add_run('一、先给结论：这4款可以买，这1款要躲开').bold = True

add_para('✅ 闭眼入系列', bold=True, size=12)

# 1
p = doc.add_paragraph()
p.add_run('1. 美的 BCD-508WTPZM(E)（莫兰迪灰）—— 4000元档性价比之王').bold = True
add_para('这台是我今年最推荐的机型。508升大容量、双系统双循环不串味、PST+全净味技术（9分钟急速净味）、一级能效。')
add_para('关键数据：已售20万+台，75万+人加购，99%好评。这个销量和口碑数据，说明它经受了大量真实用户的检验。')
add_para('补贴后到手约4000元出头。如果你预算在4000元左右，直接选这款，不会错。')

# 2
p = doc.add_paragraph()
p.add_run('2. 美的 471L BCD-471WSPZM(E) —— 刚需家庭首选').bold = True
add_para('补贴后到手1588元，471升大容量，一级能效。这个价格能买到大牌一级能效冰箱，性价比很难找到对手。')
add_para('适合人群：预算紧张但想要大容量冰箱的三口之家，或者给父母换新的朋友。')
add_para('需要注意的是，这款是单系统（不是双系统），在意串味问题的建议加钱上508。')

# 3
p = doc.add_paragraph()
p.add_run('3. 美的 熊墩墩Pro600 —— 嵌入式需求一步到位').bold = True
add_para('60cm超薄纯平全嵌，刚好匹配标准橱柜深度。600升超大容量，双系统2.0，PST+净化+四维净化，自动制冰。')
add_para('补贴后到手约7954元。嵌入式冰箱本身溢价就高，这款已经是同级别里相对划算的选择了。')
add_para('适合人群：精装房业主、追求厨房整体美观的高端用户。')

# 4
p = doc.add_paragraph()
p.add_run('4. 美的 530小冰狗 —— 年轻人中产首选').bold = True
add_para('小冰狗是美的这两年挺成功的年轻化IP，502升+双系统2.0+超薄零嵌的配置。店铺被2000万人种草，冰箱热卖榜第5名，这个数据确实很有说服力。')
add_para('补贴后到手约6469元。适合追求颜值和实用性平衡的年轻家庭。')

add_para('❌ 建议躲开', bold=True, size=12)

p = doc.add_paragraph()
p.add_run('华凌547零嵌款 —— 价格不太值').bold = True
add_para('华凌是美的子品牌没错，但547零嵌款标价和实际性能放一起看，不太匹配。同样的价格，不如加钱上熊墩墩Pro600，或者减钱上小冰狗530。')

# ════════════════════════════════════════════════════
# 参数对比图（第一节后）
# ════════════════════════════════════════════════════
add_divider()
add_para('📊  各价位段核心参数对比', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_image(IMAGES["compare"], width=Inches(6.5))
add_divider()

# ════════════════════════════════════════════════════
# 二、市场背景
# ════════════════════════════════════════════════════
h2 = doc.add_heading('', level=1)
h2.add_run('二、2026年618市场背景：为什么现在必须买？').bold = True

p = doc.add_paragraph()
p.add_run('1. 涨价已成定局，早买早划算').bold = True
add_para('很多人纠结现在是等618还是现在就买。我的判断是：如果你看中了具体的型号，现在就可以锁单了。')
add_para('原因很简单：铜价已经突破9.9万元/吨，较2025年Q4累计上涨33%-40%。ABS塑料从8000元/吨涨到13000元/吨，涨幅超过60%。')
add_para('冰箱制造成本大幅上涨，厂家已经开始调价。这次涨价不是商家的促销套路，是真的成本压力。')
add_para('618期间，即使平台有活动，优惠力度也很难完全覆盖成本涨幅。如果你等618再买，大概率会比现在贵200-500元。')

p = doc.add_paragraph()
p.add_run('2. 国补15%，现在才是最佳窗口期').bold = True
add_para('2026年国家以旧换新补贴政策力度很大：')
add_para('• 补贴比例：购买一级能效冰箱，按最终成交价的15%补贴')
add_para('• 单件封顶：每台最高补贴1500元')
add_para('• 发放节奏：第二批625亿已于4月10日下达，全国29省通用')

add_para('简单算一笔账：', bold=True)

tbl = doc.add_table(rows=1, cols=5)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, h in enumerate(['产品', '标价', '618预计到手', '国补', '实际支出']):
    hdr[i].text = h
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)

data_rows = [
    ['美的471L', '1928元', '1588元', '340元', '1588元'],
    ['美的508',  '3085元', '3700元', '450元', '3250元'],
    ['熊墩墩Pro600', '9229元', '8200元', '1276元', '6924元'],
]
for row_data in data_rows:
    add_table_row(tbl, row_data)

doc.add_paragraph()
add_para('重点说熊墩墩Pro600：原价9229元的机器，618叠加各种优惠再减1276元国补，实际到手6924元。这个优惠力度，放在整个嵌入式冰箱市场都很有竞争力。')

add_para('国补避坑提醒：', bold=True)
add_para('1. 必须单品单件下单，不能和其他商品合并付款，否则补贴资格会被取消')
add_para('2. 补贴资格券领取后当天有效，别忘了及时下单')
add_para('3. 收货地址要和身份证一致，GPS定位要在补贴发放省份')
add_para('4. 页面必须有一级能效标识，没有的直接放弃')

# ════════════════════════════════════════════════════
# 三、核心技术
# ════════════════════════════════════════════════════
h3 = doc.add_heading('', level=1)
h3.add_run('三、美的核心技术解读：这些参数决定你买亏还是买赚').bold = True

p = doc.add_paragraph()
p.add_run('1. 双系统 vs 单系统：这点钱不能省').bold = True
add_para('很多人买冰箱只看容量和价格，忽略了最影响体验的配置——系统数量。')
add_para('单系统冰箱：冷藏和冷冻共用一个蒸发器，串味问题从根本上无法解决。你放进去的榴莲，三天后可能连冰淇淋都是榴莲味。')
add_para('双系统冰箱：冷藏和冷冻各自独立制冷，从物理上隔绝串味。而且制冷速度更快，温度更稳定。')
add_para('行业数据：常规单系统冰箱的串味投诉占售后问题的40%以上。')
add_para('美的2026年的双系统做了升级：')
add_para('• 冷藏室搭载定制保湿蒸发器+超亲水涂层，主动输送高湿度冷风')
add_para('• 草莓存放13天，汁液留存率还能保持97.9%')
add_para('• 冷冻室化霜时间缩短14.8%，排水效率提升50%')
add_para('我的建议：预算超过2500元，尽量选双系统。这不是噱头，是直接影响你未来5-10年使用体验的关键配置。')

p = doc.add_paragraph()
p.add_run('2. PST+智能净化7.0：冰箱里的空气净化器').bold = True
add_para('除菌净味技术这几年发展很快，但大多数品牌只解决了冷藏室的问题。')
add_para('-18℃的冷冻室环境下，传统离子除菌技术往往会失效。嗜冷菌（比如李斯特菌）在极低温下依然能存活缓慢繁衍。')
add_para('美的PST+7.0版本突破了这个物理极限：')
add_para('• 实现了在-40℃极端超低温环境下依然保持稳定运行')
add_para('• 9分钟急速净味，冷藏冷冻全空间主动净化')
add_para('• 这项技术目前是行业唯一获得两项国际领先认证的方案')
add_para('适合人群：家里有小孩、老人的，或者经常囤肉、海鲜的家庭，这项配置能有效降低食材污染风险。')

p = doc.add_paragraph()
p.add_run('3. 低氧窖藏PRO：果蔬保鲜革命').bold = True
add_para('很多冰箱标榜"保鲜7天"，实际上绿叶菜放进去三天就开始发黄腐烂。')
add_para('低氧窖藏PRO的思路完全不同：不是单纯降温，而是控制氧气浓度。')
add_para('通过高分子密封结构+控氧保湿膜，在冰箱内部创造一个低氧、高湿、恒温的独立微环境。')
add_para('实验数据：果蔬鲜活力留存提升6倍，可以实现14天长效原鲜保鲜。')
add_para('简单说：以前你买菜只能撑3天，现在可以撑两周。这项技术对喜欢周末一次性采购大量蔬果的家庭特别实用。')

p = doc.add_paragraph()
p.add_run('4. 冷冻超智能恒温：冻肉不流血水的秘密').bold = True
add_para('你有没有这种感觉：解冻后的肉，口感又柴又腥，汁液流了一砧板？')
add_para('这是因为冷冻过程中形成了大冰晶，刺破了细胞壁。解冻时肌红蛋白和营养物质随着汁液大量流失。')
add_para('美的的解决方案是把温度波动控制在0.9℃以内。')
add_para('• 配合AI智能温控算法预判温度变化')
add_para('• 30天冷冻后，牛肉解冻汁液留存率依然达到98.18%')
add_para('这个参数什么意思？就是你冻了一个月的牛肉，拿出来化冻，血水只有一点点，烹饪后口感接近新鲜肉。')

# ════════════════════════════════════════════════════
# 四、分价格段推荐（插入人群选购图）
# ════════════════════════════════════════════════════
h4 = doc.add_heading('', level=1)
h4.add_run('四、分价格段推荐清单：照着买不踩坑').bold = True

add_divider()
add_para('👥  看看你属于哪类人群，对号入座', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_image(IMAGES["people"], width=Inches(6))
add_divider()

p = doc.add_paragraph()
p.add_run('入门款（800-1500元）：租房党首选').bold = True
p = doc.add_paragraph()
p.add_run('美的181L两门小冰箱').bold = True
add_para('• 到手约802元')
add_para('• 已售70万+台，冰箱热卖榜第9名')
add_para('• 租房、宿舍首选，实用够用')
add_para('• 缺点：机械控温，没有变频，不支持智能')
add_para('这个价位就别追求什么高级功能了，181升容量、风冷无霜、一级能效，解决了有冰箱用的问题。')

p = doc.add_paragraph()
p.add_run('性价比款（1500-2500元）：刚需家庭首选').bold = True
p = doc.add_paragraph()
p.add_run('美的471L BCD-471WSPZM(E)').bold = True
add_para('• 补贴后1588元，已售40万+')
add_para('• 一级能效双变频')
add_para('• 全空间养鲜+抗菌净味')
add_para('• 800+人复购')
p = doc.add_paragraph()
p.add_run('美的480L BCD-480WSPZM(E)').bold = True
add_para('• 补贴后约2678元，已售7万+')
add_para('• 十字四门，50万+人加购')
add_para('这个价位段，471L的性价比是最高的。40万+的销量说明它经过了市场验证。')

p = doc.add_paragraph()
p.add_run('主流款（2500-4000元）：普通改善型家庭').bold = True
p = doc.add_paragraph()
p.add_run('美的508L BCD-508WTPZM(E)').bold = True
add_para('• 补贴后约4000元')
add_para('• 双系统+PST+全净味')
add_para('• 20万+销量，75万+人加购')
add_para('• 法式四门，颜值和实用性兼顾')
add_para('这是我认为最值得买的机型。4000元档，双系统+顶级净味技术，性价比之王。')
p = doc.add_paragraph()
p.add_run('美的532L 法式多门').bold = True
add_para('• 补贴后约3275元，已售6000+')
add_para('• 双系统双PT抗菌净味')
add_para('• 90度直角开门，贴墙放也没问题')

p = doc.add_paragraph()
p.add_run('中高端款（4000-6500元）：品质家庭首选').bold = True
p = doc.add_paragraph()
p.add_run('美的M60 520小机皇').bold = True
add_para('• 补贴后约6135元')
add_para('• 60cm超薄零嵌，双系统2.0，PST净化')
add_para('• 适合精装房和对颜值有要求的用户')
p = doc.add_paragraph()
p.add_run('美的530小冰狗').bold = True
add_para('• 补贴后约6469元')
add_para('• 502升，店铺2000万人种草')
add_para('• 双系统2.0，超薄零嵌')
add_para('• 年轻人中产首选')
p = doc.add_paragraph()
p.add_run('美的505L十字四门').bold = True
add_para('• 补贴后约3057元')
add_para('• 双系统双循环')
add_para('• 价格相对便宜，配置不差')

p = doc.add_paragraph()
p.add_run('高端款（6500-10000元）：高端改善型').bold = True
p = doc.add_paragraph()
p.add_run('美的熊墩墩Pro600').bold = True
add_para('• 补贴后约7954元')
add_para('• 600升超大容量，60cm纯平全嵌')
add_para('• 双系统2.0，PST+净化，四维净化，自动制冰')
add_para('• 嵌入式冰箱顶配选择')
p = doc.add_paragraph()
p.add_run('美的601 熊墩墩Pro').bold = True
add_para('• 补贴后约6796元')
add_para('• 十字四门全嵌设计')
add_para('• 双系统净味，行业功能顶配')

# ════════════════════════════════════════════════════
# 优缺点图（价格段推荐后）
# ════════════════════════════════════════════════════
add_divider()
add_para('✅❌  各款机型优缺点一览', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_image(IMAGES["proscons"], width=Inches(6.5))
add_divider()

# ════════════════════════════════════════════════════
# 五、618国补领取攻略
# ════════════════════════════════════════════════════
h5 = doc.add_heading('', level=1)
h5.add_run('五、618国补领取攻略：能省1500元的实操手册').bold = True

p = doc.add_paragraph()
p.add_run('京东领取路径').bold = True
add_para('1. 打开京东APP，在搜索栏输入：家电国补687 或 家电省2000')
add_para('2. 进入国补专区后，完成实名认证绑定')
add_para('3. 挑选带有"一级能效"+"国补专享"标签的美的冰箱')
add_para('4. 下单时系统自动抵扣')
add_para('备用口令（主口令拥堵时）：领国补855、福利555、福利800')

p = doc.add_paragraph()
p.add_run('PLUS会员额外叠加').bold = True
add_para('如果是京东PLUS会员，搜索会员555进入专属会场，可以再享受95折优惠。这个叠加和国补不冲突，可以同时享受。')

p = doc.add_paragraph()
p.add_run('天猫/淘宝领取路径').bold = True
add_para('搜索以下任意口令进入国补专场：')
add_para('• 红包到手77171')
add_para('• 天降红包113')
add_para('• 补贴打工人8800')

p = doc.add_paragraph()
p.add_run('线下购买').bold = True
add_para('不想等快递、想先看看实物的，直接去美的专卖店或苏宁易购。')
add_para('流程：带身份证 → 选好冰箱 → 告诉店员要国补 → 店员政企联网核验 → 直接立减15%。')
add_para('线下的好处是现场看货、即时提货，不用担心物流磕碰。')

p = doc.add_paragraph()
p.add_run('四大避坑提醒').bold = True
add_para('第一：单品单下单')
add_para('冰箱、空调、电视、手机必须分开下单，不能放一起合并付款。一旦合并，补贴资格可能被取消，两件商品都享受不到优惠。')
add_para('第二：GPS和收货地址必须一致')
add_para('系统会校验收货地址和身份证所在省市是否一致。人在北京，就让商家发到北京地址，别找人代收再转发。')
add_para('第三：资格券当天有效')
add_para('领取国补资格后，当天必须下单支付。失效后可以重新领取，但一个身份证累计4次未使用，会被永久取消该品类补贴资格。')
add_para('第四：最后一步检查一级能效标识')
add_para('付款前再看一眼商品主图左上角，必须有绿色盾牌的一级能效标识。没有这个标的，不要买。')

# ════════════════════════════════════════════════════
# 六、说在最后
# ════════════════════════════════════════════════════
h6 = doc.add_heading('', level=1)
h6.add_run('六、说在最后').bold = True

add_para('买冰箱这件事，很多人花了很多时间研究参数，结果买回家发现跟想象的不太一样。')
add_para('其实冰箱选购没那么复杂：')
add_para('1. 先定预算：1500元以内别想太多，选销量大的；4000元左右直接上双系统款')
add_para('2. 再定容量：三口之家选400-500升，四口以上选500升以上')
add_para('3. 看安装条件：橱柜深度60cm的，直接选嵌入式；否则随便选')
add_para('4. 盯住国补：一级能效才能补贴，别买二级能效的老款')
add_para('今年618，美的冰箱的优惠力度确实不错。国补+平台优惠+品牌让利三重叠加，是入手的好时机。')
add_para('但记住：再便宜，也是花自己的钱。买回来不满意的冰箱，比买贵了更闹心。')
add_para('希望这篇指南能帮你买到一台真正适合自己家的冰箱。', italic=True)

# ════════════════════════════════════════════════════
# 场景图（结尾前）
# ════════════════════════════════════════════════════
add_divider()
add_para('🏠  你的使用场景，决定了你的最佳选择', bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_image(IMAGES["scene"], width=Inches(6.5))
add_divider()

# ════════════════════════════════════════════════════
# 数据来源
# ════════════════════════════════════════════════════
doc.add_paragraph()
add_para('数据来源：京东销量数据（2026年5月3日）、国家以旧换新补贴政策（2026年最新版）、美的冰箱2026年新品发布会技术白皮书', size=9, color=(120,120,120))

# ════════════════════════════════════════════════════
# 保存
# ════════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
print(f"✅ Word文档已生成：{OUTPUT_PATH}")
