#!/usr/bin/env python3
"""Humanizer + Word Export Pipeline"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUTS_DIR = "/Users/ltn/Downloads/ARTICLE WRITER/outputs"
SOURCE_FILE = f"{OUTPUTS_DIR}/T6_final_头条.md"
OUTPUT_DOCX = f"{OUTPUTS_DIR}/高端大屏电视_国产vs索尼.docx"

# =====================
# 去AI味处理
# =====================
def humanize_text(text):
    """深度去AI味处理"""
    replacements = [
        # 消除"首先/其次/最后"模板
        (r"首先，", "先说"),
        (r"其次，", "再说"),
        (r"最后，", "再补充一点"),
        (r"第一，", "一方面"),
        (r"第二，", "另一方面"),
        (r"第三，", "再一方面"),
        # 书面词→口语
        (r"然而", "但"),
        (r"因此", "所以"),
        (r"由此可见", "说白了"),
        (r"综上所述", "总的来说"),
        (r"值得注意的是", "要我说"),
        (r"从\s*技术\s*角度\s*说", "技术层面来看"),
        (r"从\s*纯\s*技术", "单说技术"),
        # 消除"极端/非常"滥用
        (r"极端\s*能\s*打", "相当能打"),
        (r"非常\s*能\s*打", "相当能打"),
        # 承认边界感
        (r"已经是\s*天花板\s*级别", "已经是很高的水平"),
        (r"已经\s*做到\s*超越", "已经在逼近"),
        # 减少绝对化
        (r"差距\s*正在\s*缩小", "差距在慢慢缩小"),
        # 场景代入感
        (r"比如说", "就像"),
        (r"举个例子", "说个场景"),
        # 消除"这套/这个"过多次使用
        (r"索尼那套", "索尼那套"),
        (r"这套系统", "这套方案"),
    ]

    for old, new in replacements:
        text = re.sub(old, new, text)

    # 打散长句：超过30字的长句尝试拆分
    sentences = re.split(r'([。！？])', text)
    result = []
    for i in range(0, len(sentences)-1, 2):
        s = sentences[i]
        punct = sentences[i+1] if i+1 < len(sentences) else ''
        if len(s) > 50 and '，' in s:
            # 拆分超长句
            parts = s.split('，')
            new_parts = []
            for j, part in enumerate(parts):
                if j == 0:
                    new_parts.append(part)
                elif len(part) > 20:
                    new_parts.append('，' + part)
                else:
                    new_parts.append(part)
            s = ''.join(new_parts)
        result.append(s + punct)

    return ''.join(result)

# =====================
# 解析 Markdown
# =====================
def parse_markdown(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    blocks = []
    current_text = []

    for line in content.split('\n'):
        # 图片处理
        if line.startswith('!['):
            # 先输出当前文本段落
            if current_text:
                blocks.append(('text', '\n'.join(current_text)))
                current_text = []
            # 提取图片路径
            match = re.search(r'\]\(([^)]+)\)', line)
            if match:
                img_path = match.group(1)
                blocks.append(('image', img_path))
        # 标题
        elif line.startswith('# '):
            if current_text:
                blocks.append(('text', '\n'.join(current_text)))
                current_text = []
            blocks.append(('h1', line[2:]))
        elif line.startswith('## '):
            if current_text:
                blocks.append(('text', '\n'.join(current_text)))
                current_text = []
            blocks.append(('h2', line[3:]))
        elif line.startswith('**'):
            # 加粗段落，可能是小标题
            if current_text:
                blocks.append(('text', '\n'.join(current_text)))
                current_text = []
            blocks.append(('subhead', re.sub(r'\*\*', '', line)))
        elif line.startswith('---'):
            blocks.append(('hr', ''))
        elif line.startswith('*') and line.endswith('*'):
            # 注释行，跳过
            continue
        else:
            current_text.append(line)

    if current_text:
        blocks.append(('text', '\n'.join(current_text)))

    return blocks

# =====================
# 生成 Word
# =====================
def create_word(blocks, img_dir, output_path):
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for block_type, block_content in blocks:
        if block_type == 'h1':
            p = doc.add_heading(block_content, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block_type == 'h2':
            p = doc.add_heading(block_content, level=2)
        elif block_type == 'subhead':
            p = doc.add_paragraph()
            run = p.add_run(block_content)
            run.bold = True
            run.font.size = Pt(12)
        elif block_type == 'hr':
            p = doc.add_paragraph()
            p_fmt = p.paragraph_format
            p_fmt.space_before = Pt(6)
            p_fmt.space_after = Pt(6)
            # 添加横线
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif block_type == 'image':
            img_full_path = os.path.join(img_dir, os.path.basename(block_content))
            if os.path.exists(img_full_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # 宽度设为6英寸，保持比例
                run.add_picture(img_full_path, width=Inches(6.0))
            else:
                p = doc.add_paragraph(f'[图片: {block_content} 未找到]')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block_type == 'text':
            # 逐段处理
            for para in block_content.split('\n\n'):
                para = para.strip()
                if not para:
                    continue
                # 加粗处理
                p = doc.add_paragraph()
                # 处理 **bold** 文本
                parts = re.split(r'(\*\*[^*]+\*\*)', para)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                # 段后距
                p.paragraph_format.space_after = Pt(8)

    doc.save(output_path)
    print(f"Word 文档已生成: {output_path}")

# =====================
# 主流程
# =====================
if __name__ == "__main__":
    # 1. 读取原始文章
    with open(SOURCE_FILE, 'r', encoding='utf-8') as f:
        original = f.read()

    # 2. 去AI味处理
    humanized = humanize_text(original)

    # 3. 保存去AI味版本
    humanized_path = f"{OUTPUTS_DIR}/T6_Humanized.md"
    with open(humanized_path, 'w', encoding='utf-8') as f:
        f.write(humanized)
    print(f"去AI味版本: {humanized_path}")

    # 4. 解析并生成Word
    blocks = parse_markdown(humanized_path)
    create_word(blocks, OUTPUTS_DIR, OUTPUT_DOCX)
    print(f"Word文档: {OUTPUT_DOCX}")
