#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""build_article_docx.py — Super Writer 终稿 md + 配图 → 图文 Word。
固化自 2026-06-17 电视稿（python-docx）。零网络、单文件、可复用于空调/热水器等任何图文稿。

用法:
  python3 build_article_docx.py \
    --md  outputs/03_风控主编终稿_电视.md \
    --out outputs/TCL电视_图文稿.docx \
    --cover outputs/封面.png \
    --img1  outputs/T5_面板三件套档位图.png \
    --img2  outputs/T5_四维选购对比.png

规则:
  - 主标题(# )居中大字；其后的第一段正文前自动插入封面图。
  - 整行含「信息图1」→ 插 img1；含「信息图2」→ 插 img2（关键词可 --img1-kw/--img2-kw 覆盖）。
  - 默认用于排版预览：图路径缺失或文件不存在 → 留红字占位。
  - 正式交付前必须先运行 Scripts/validate_visual_outputs.py，Word 只应嵌入已落盘、已验收图片。
  - 列表项(- )、加粗(**)、二级标题(## )、正文首行缩进 均自动处理。
"""
import argparse, re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CN_BODY, CN_HEAD = "宋体", "黑体"
IMG_W = 15.5  # content width cm on A4 with 2.5cm side margins


def set_font(run, size=11, bold=False, color=None, name_cn=CN_BODY):
    run.font.size = Pt(size); run.font.bold = bold; run.font.name = name_cn
    rPr = run._element.get_or_add_rPr(); rPr.get_or_add_rFonts().set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_runs_bold(para, text, size=11, name_cn=CN_BODY, base_bold=False):
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2]); set_font(r, size, True, None, name_cn)
        else:
            r = para.add_run(part); set_font(r, size, base_bold, None, name_cn)


def add_pic(doc, path, w=IMG_W):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Cm(w))


def add_placeholder(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); set_font(r, 10, True, (0xC0, 0x00, 0x00))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--md', required=True)
    ap.add_argument('--out', required=True)
    ap.add_argument('--cover')
    ap.add_argument('--img1')
    ap.add_argument('--img2')
    ap.add_argument('--img1-kw', default='信息图1')
    ap.add_argument('--img2-kw', default='信息图2')
    ap.add_argument('--strict-images', action='store_true', help='fail if any requested image is missing')
    a = ap.parse_args()

    doc = Document(); sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.2)
    sec.left_margin = sec.right_margin = Cm(2.5)

    def maybe_pic(path):
        if path and os.path.exists(path):
            add_pic(doc, path)
        else:
            if a.strict_images:
                raise FileNotFoundError(f'正式交付图片缺失：{path or "未指定"}')
            add_placeholder(doc, f'【此处插图缺失：{path or "未指定"}】')

    with open(a.md, encoding='utf-8') as f:
        lines = f.read().splitlines()

    need_cover = a.cover is not None  # cover 待补时不插封面
    for line in lines:
        s = line.strip()
        if not s:
            continue
        # doc title
        if s.startswith('# ') and not s.startswith('## '):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            add_runs_bold(p, s[2:].strip(), 18, CN_HEAD, True)
            continue
        # drop cover right before first body content
        if need_cover:
            maybe_pic(a.cover); need_cover = False
        # section heading
        if s.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
            add_runs_bold(p, s[3:].strip(), 13, CN_HEAD, True)
        elif a.img1_kw in s or '信息图1' in s:
            maybe_pic(a.img1)
        elif a.img2_kw in s or '信息图2' in s:
            maybe_pic(a.img2)
        elif s.startswith('- '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.5
            r = p.add_run('· '); set_font(r, 11)
            add_runs_bold(p, s[2:].strip(), 11)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(0.74)
            add_runs_bold(p, s, 11)

    doc.save(a.out); print('saved:', a.out)


if __name__ == '__main__':
    main()
