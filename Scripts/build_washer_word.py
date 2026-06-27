#!/usr/bin/env python3
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFilter


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
MD = OUT / "06_去AI味优化终稿_洗衣机.md"
DOCX = OUT / "洗衣机去AI味图文稿_20260617.docx"

FONT_CN = "PingFang SC"
FONT_FALLBACK = "Arial"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_FALLBACK
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_para_spacing(paragraph, before=0, after=7, line=1.18):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_border(paragraph, color="D9E2EC", size="8", space="8"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def prepare_cover():
    src = OUT / "source_docx_media" / "image2.jpeg"
    if not src.exists():
        return OUT / "封面_洗衣机_滚筒更省心.png"

    img = Image.open(src).convert("RGB")
    # Crop away the control-panel/logo strip for a cleaner neutral article header.
    w, h = img.size
    crop = img.crop((0, 145, w, min(h, 820)))
    target_w, target_h = 1400, 620
    canvas = Image.new("RGB", (target_w, target_h), (245, 248, 250))
    scale = max(target_w / crop.width, target_h / crop.height)
    resized = crop.resize((int(crop.width * scale), int(crop.height * scale)))
    left = (target_w - resized.width) // 2
    top = (target_h - resized.height) // 2
    canvas.paste(resized, (left, top))
    # Subtle left-to-right white wash so headline could be overlaid later if needed.
    overlay = Image.new("RGBA", canvas.size, (255, 255, 255, 0))
    od = ImageDraw.Draw(overlay)
    for x in range(0, 620):
        alpha = int(130 * (1 - x / 620))
        od.line((x, 0, x, target_h), fill=(255, 255, 255, alpha))
    canvas = Image.alpha_composite(canvas.convert("RGBA"), overlay).convert("RGB")
    out = OUT / "封面_洗衣机_真实场景_临时版.jpg"
    canvas.save(out, quality=92)
    return out


def parse_markdown(text):
    blocks = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            blocks.append(("blank", ""))
        elif line.startswith("# "):
            blocks.append(("title", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        elif line.startswith("!["):
            m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
            if m:
                blocks.append(("image", (m.group(1), Path(m.group(2)))))
        else:
            blocks.append(("p", line))
    return blocks


def add_image(doc, path, width=5.85):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    set_para_spacing(p, before=6, after=10)


def build():
    cover = prepare_cover()
    text = MD.read_text(encoding="utf-8")
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_FALLBACK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11.2)
    normal.font.color.rgb = RGBColor(34, 42, 52)

    for block_type, value in parse_markdown(text):
        if block_type == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, size=22, bold=True, color=(16, 32, 48))
            set_para_spacing(p, before=0, after=10, line=1.12)
            add_bottom_border(p)
            add_image(doc, cover, width=6.15)
        elif block_type == "h2":
            p = doc.add_paragraph()
            run = p.add_run(value)
            set_run_font(run, size=15, bold=True, color=(15, 78, 120))
            set_para_spacing(p, before=13, after=6, line=1.12)
        elif block_type == "p":
            p = doc.add_paragraph()
            run = p.add_run(value)
            set_run_font(run, size=11.2, color=(34, 42, 52))
            p.paragraph_format.first_line_indent = Pt(22)
            set_para_spacing(p, before=0, after=7, line=1.18)
        elif block_type == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(value)
            set_run_font(run, size=11.2, color=(34, 42, 52))
            set_para_spacing(p, before=0, after=4, line=1.15)
        elif block_type == "image":
            alt, path = value
            add_image(doc, path, width=3.55)
        elif block_type == "blank":
            continue

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("牛科技说 · 发布稿")
    set_run_font(r, size=9, color=(120, 132, 145))

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
