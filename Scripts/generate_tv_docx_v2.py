#!/usr/bin/env python3
"""Generate Word doc from revised TV article (v2 with TCL T7M Pro / Q9M Pro)."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUTS_DIR = os.path.join(BASE_DIR, "outputs")


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return h


def add_para(doc, text, bold=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p


def add_image_with_caption(doc, img_path, caption, max_width=6.0):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(max_width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r.italic = True
    else:
        add_para(doc, f"[图片缺失: {os.path.basename(img_path)}]", size=10, color=RGBColor(0xcc, 0x00, 0x00))


def add_table_from_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        from docx.oxml import OxmlElement
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '2B579A')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9)
            if ri % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F2F7FB')
                shading.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()


def build_doc():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing = 1.5

    # Title
    title = doc.add_heading('海信、TCL、创维电视怎么选？三分钟看懂，别再交智商税', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        run.font.size = Pt(22)

    # Cover image
    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_封面图.png'),
        '2026年618海信、TCL、创维三大品牌电视选购横评')

    # Opening paragraph
    add_para(doc, '买电视这事儿，一句话就能说清楚：要画质选海信，要性价比和游戏体验选TCL，要颜值和护眼选创维。2025年中国电视零售量跌到2763万台（奥维睿沃数据），近十年新低，但Mini LED占比已经突破80%（洛图科技数据），"量减质升"的趋势摆在那儿。今天拿数据说话，把三大品牌各自的路数掰扯明白。', space_after=12)

    # Section 1: Market overview
    add_heading(doc, '电视市场格局：四家国产品牌吃掉80%以上份额', level=2)

    add_para(doc, '2025年全球彩电出货量大约2.64亿台（奥维睿沃+新浪财经数据），同比基本持平。中国市场这边就比较刺激了——零售量2763万台，出货量3289.5万台（洛图科技），两个数字口径不同，零售量是消费者实际掏钱买的，出货量是厂商发给渠道的，出货大于零售很正常。但格局已经定了：海信、TCL、小米、创维四家吞了超过80%的份额，前八大品牌占96%。长虹、海尔、康佳三家出货量加一块儿才103万台，同比还掉了12.2%，第二阵营在加速萎缩。三星、LG在中国市场更是持续边缘化，基本在慢慢退出家电业务（新浪财经报道）。')

    add_para(doc, '不过总量跌不代表没机会。大屏化趋势很明显，98寸和100寸的销量份额已经到2.8%左右，同比翻倍。换机周期拉长到10年以上（奥维睿沃数据），中国百户保有量超过120台，市场高度饱和。这种存量博弈下，Mini LED在2025年618期间占比突破80%（洛图科技数据），从"高端选项"变成了"标配"。')

    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_技术路线对比图.png'),
        '三大品牌核心技术路线对比（数据来源：品牌官方及京东商品页，2026年5月）')

    # Section 2: Hisense
    add_heading(doc, '海信：RGB-Mini LED画质天花板路线', level=2)

    add_para(doc, '先看海信的财务底盘。2025年年报显示，海信视像营收576.79亿元，同比微降1.45%，但归母净利润24.54亿元，反而涨了9.24%。营收降利润升，摆明了海信在主动调结构——宁可少卖也不打价格战，走"利润优先"。')

    add_para(doc, '技术层面，海信的杀手锏是RGB-Mini LED。旗舰UX 2026款搞了个"玲珑4芯真彩背光"，在红、绿之外加了颗青色（Cyan）发光芯片，填补了480-520nm波段的显示空白（量子位报道）。配合信芯AI H7 Pro画质芯片，这台机器堆到了10000尼特峰值亮度、43008个背光分区、110% BT.2020色域覆盖。京东自营旗舰店定价：85寸34999元，100寸54999元，冲着高端家庭和专业影音发烧友去的。黑曜屏Ultra反射率只有1.28%，强光下画面照样通透。帝瓦雷6.2.2声道音响系统在这个价位算领先配置，还支持2D转3D和Mac模式。')

    add_para(doc, '海信还有一条独门产品线——激光电视。Omdia数据显示（2022年统计），海信激光电视全球份额53.5%，8K激光电视属于首发量产级别。想要100寸以上的超大显示面积，又对环境光控制有条件的话，激光电视的护眼和尺寸优势值得认真考虑。')

    add_heading(doc, 'RGB-Mini LED路线的优缺点', level=3)
    for line in ['优点：', '- 画质参数领先：10000尼特+43008分区的组合，对比度和HDR表现优势明显', '- 自研画质芯片加持：信芯H7 Pro是国产品牌里少有的独立画质芯片，实时调校能力强', '- 色域覆盖广：110% BT.2020配合青色芯片创新，色彩还原度高于传统量子点方案', '', '不足：', '- 价格门槛高：旗舰85寸34999元，对大多数家庭预算压力不小', '- 激光电视受限：需要专业幕布和遮光环境，安装条件苛刻', '- 入门款竞争力一般：海信中低端性价比不如TCL同价位产品']:
        if line.startswith('优点') or line.startswith('不足'):
            add_para(doc, line, bold=True, size=11)
        elif line:
            add_para(doc, line, size=11)

    # Section 3: TCL (REVISED)
    add_heading(doc, 'TCL：SQD-Mini LED性价比+技术普惠路线', level=2)

    add_para(doc, 'TCL的财务数据很炸。2025年年报显示，TCL电子营收1145.83亿港元，同比涨15.4%，归母净利润24.95亿港元，同比暴增41.8%。增长主要靠Mini LED产品线爆发——前三季度Mini LED出货量224万台，同比增153.3%（东方财富研报），全球占比11.4%，国内占比20.6%。北美市场Mini LED增速384.5%，欧洲124.1%。东方财富研报还提到，海信+TCL在600美元以上高端市场的合计份额，从2023年的22%飙到2025年一季度的39%。')

    add_para(doc, 'TCL的底层优势是"屏厂一体化"。旗下华星光电是全球面板大厂，从面板到整机一条龙自家产线，成本控制在国产品牌里很有底气。2026年3月春季发布会上，TCL把SQD-Mini LED（超级量子点）技术从旗舰X11L下放到主流价位，一口气推出三款新品：Q9M Pro定位"画质轻旗舰"，T7M Pro定位"甜品级换代款"，T7M Ultra定位"家庭普惠款"（据IT之家、中华网报道）。')

    add_para(doc, '代表产品T7M Pro 85寸，SQD-Mini LED加持，搭载超级蝶翼星曜屏，最高1152个万象分区（据IT之家），绚彩XDR 2200尼特峰值亮度，100% BT.2020全局高色域。4K 150Hz原生高刷可开启300Hz超高刷，安桥2.1.2 Hi-Fi音响，四路满血HDMI 2.1接口。灵控系统3.0明确无开机广告，TSR AI光色同控芯片支持云端+终端双重画质优化。85寸零售价9999元（据慢慢买比价网），国补后约8499元。这个价位拿到SQD-Mini LED+万象分区+100% BT.2020，85寸市场里性价比很能打。')

    add_para(doc, '更进一步的Q9M Pro 85寸，定位"画质轻旗舰"。绚彩XDR峰值亮度拉到5000尼特，万象分区数更高（98寸版本达3552个，据雷科技），超级蝶翼星曜屏的原生对比度7000:1。85寸国补价10499元（据雷科技），比T7M Pro贵了约2000元，但亮度和分区数明显上一个台阶。98寸国补价18499元（据雷科技），是98寸Mini LED市场里有竞争力的选择。配置上同样是4K 150Hz→300Hz、安桥2.1.2 Hi-Fi、灵控系统3.0无开机广告。')

    add_heading(doc, 'SQD-Mini LED路线的优缺点', level=3)
    for line in ['优点：', '- 技术普惠力度大：SQD-Mini LED从旗舰X11L下放到6000元档（据IT之家），同价位多数竞品还在用普通Mini LED', '- 性价比突出：T7M Pro 85寸国补后约8499元，Q9M Pro 85寸国补后10499元，参数组合在各自价位段有竞争力', '- 游戏体验好：300Hz超高刷+四路满血HDMI 2.1，PS5和Xbox适配成熟，无开机广告口碑不错', '- 全产业链优势：华星光电面板自研自产，成本控制和技术迭代速度有保障', '', '不足：', '- 画质芯片依赖通用方案：没有类似海信信芯的独立自研画质芯片，TSR AI芯片偏重算法优化而非硬件级调校', '- 高端线差异化不够：旗舰X11L在画质天花板层面跟海信UX仍有差距', '- 品牌溢价低：高端市场的品牌认知度不如海信']:
        if line.startswith('优点') or line.startswith('不足'):
            add_para(doc, line, bold=True, size=11)
        elif line:
            add_para(doc, line, size=11)

    # Section 4: Skyworth
    add_heading(doc, '创维：壁纸电视美学+护眼路线', level=2)

    add_para(doc, '创维的财务状况有点意思。2025年年报显示，创维集团营收703.24亿元，同比涨了8.2%，但归母净利润只有3.56亿元，反而掉了37.3%。利润下滑主要跟新能源等第二曲线的投入有关，电视主业整体处于战略转型期。')

    add_para(doc, '创维的核心策略很聪明——避开跟海信、TCL的参数军备竞赛，走"家居美学+护眼"的差异化路线。代表产品A7H Pro 85寸，京东商品页标价9763.52元（2026年5月京东自营价，含国补后约8299元）。这台机器的核心卖点不在参数堆砌，而在设计感——约3-4cm超薄壁纸设计，贴墙安装后跟挂画差不多。画境屏用AG/LR低反技术，99% DCI-P3色域，ΔE≤0.6的专业级色准，1344个Mini LED分区，3800尼特峰值亮度（2026年85寸电视性价比排行榜，新浪财经）。哈曼卡顿双Soundbar音响加分不少，莱茵TÜV低蓝光和无频闪双认证对于有孩家庭来说实打实的健康考量。4K 150Hz刷新率支持倍频300Hz，日常追剧和看球够用。')

    add_para(doc, '旗舰款A10H 100寸定价25599元（京东商品页），6120分区SQD Mini LED，跟京东方合作的面板方案，极黑广角类纸屏。酷开AIOS系统接入了DeepSeek-R1大模型（京东商品页），智能交互上有不错的创新。')

    add_heading(doc, '壁纸Mini LED路线的优缺点', level=3)
    for line in ['优点：', '- 家居美学突出：3-4cm超薄壁纸设计是国产品牌里少有的"去家电化"产品，适合讲究客厅整体风格的家庭', '- 护眼配置到位：莱茵TÜV双认证+ACD零频闪，有孩家庭的健康考量比较充分', '- 色准专业级：ΔE≤0.6在万元价位段属于领先水平，对色彩敏感的用户会比较满意', '', '不足：', '- 性价比不占优：9000元上下的价位在85寸市场面临TCL和海信的强力竞争', '- 画质天花板不够高：3800尼特和1344分区在中高端算主流水平，跟海信UX和TCL Q9M Pro有明显差距', '- 品牌盈利承压：净利润掉了37.3%，研发投入可能受限，长期产品力存在不确定性']:
        if line.startswith('优点') or line.startswith('不足'):
            add_para(doc, line, bold=True, size=11)
        elif line:
            add_para(doc, line, size=11)

    # Section 5: Comparison table
    add_heading(doc, '85寸主力机型参数横评', level=2)

    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_参数对比图.png'),
        '85寸主力机型核心参数对比（数据来源：京东商品页、IT之家、雷科技、2026年85寸电视性价比排行榜）')

    add_para(doc, '四款85寸主力机型放一起看看：', bold=True)

    headers = ['参数', '创维 A7H Pro', 'TCL T7M Pro', 'TCL Q9M Pro', '海信 85E3Q']
    rows = [
        ['参考价', '¥9,763\n（含国补约¥8,299）', '¥9,999\n（国补约¥8,499）', '国补¥10,499', '¥4,999'],
        ['背光技术', 'Mini LED\n1344分区', 'SQD-Mini LED\n万象分区', 'SQD-Mini LED\n万象分区', 'AI色彩增强'],
        ['峰值亮度', '3800尼特', '绚彩XDR\n2200尼特', '绚彩XDR\n5000尼特', '未标注'],
        ['刷新率', '4K 150Hz\n（倍频300Hz）', '4K 150Hz\n（可开300Hz）', '4K 150Hz\n（可开300Hz）', '144Hz+MEMC'],
        ['色域', '99% DCI-P3', '100% BT.2020', '100% BT.2020', '未标注'],
        ['色准', 'ΔE≤0.6', '未标注', '未标注', '未标注'],
        ['音响', '哈曼卡顿\n双Soundbar', '安桥\n2.1.2 Hi-Fi', '安桥\n2.1.2 Hi-Fi', '未标注'],
        ['特色功能', '超薄壁纸设计\n+莱茵护眼', '无开机广告\n+TSR AI芯片', '无开机广告\n+TSR AI芯片', 'AI语音\n+低蓝光护眼'],
    ]
    add_table_from_data(doc, headers, rows)

    add_para(doc, '数据来源：TCL T7M Pro参数据IT之家（2026年3月）和慢慢买比价网；Q9M Pro价格据雷科技（2026年3月），参数据IT之家和中华网；创维和海信数据据京东商品页和2026年85寸电视性价比排行榜（新浪财经）。', size=9, color=RGBColor(0x66, 0x66, 0x66))

    add_para(doc, '几个关键差异一目了然：TCL Q9M Pro在10499元国补价给了SQD-Mini LED+5000尼特+100% BT.2020，参数堆叠在万元出头价位非常有竞争力；TCL T7M Pro国补后约8499元，SQD-Mini LED+万象分区+无开机广告的配置适合预算在8000元档的用户；创维A7H Pro在色准ΔE≤0.6和超薄壁纸设计上独树一帜，对美学有要求的家庭会更倾向；海信85E3Q定价4999元，85寸入门市场的守门员。')

    # Section 6: Scenario recommendations
    add_heading(doc, '按场景选购：搞清需求比纠结品牌重要', level=2)

    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_价格段选购图.png'),
        '2026年618电视按预算对号入座（价格来源：京东、IT之家、雷科技，2026年3-5月）')

    add_para(doc, '预算5000元以内，想买85寸大屏', bold=True, size=12)
    add_para(doc, '海信85E3Q（4999元，京东商品页）是这个价位85寸的守门员。如果预算能再加500-1000元，可以考虑TCL子品牌雷鸟的85R69A（5799元，京东商品页），Mini LED 400+分区的参数在5K价位有竞争力。不过要注意雷鸟是TCL子品牌，不是TCL主品牌产品线。')

    add_para(doc, '预算8000到10000元，兼顾画质和日常', bold=True, size=12)
    add_para(doc, 'TCL T7M Pro 85寸（国补约8499元，据慢慢买）是这个价位段的有力竞争者——SQD-Mini LED+万象分区+2200尼特+100% BT.2020，加上无开机广告和安桥2.1.2音响，日常追剧和偶尔打游戏都能覆盖。预算再往上走，创维A7H Pro（含国补约8299元）画质和美学更进一步，超薄壁纸设计+哈曼卡顿的组合适合对音画品质和家居风格都有要求的家庭。')

    add_para(doc, '预算10000到20000元，追求大屏游戏或更高画质', bold=True, size=12)
    add_para(doc, 'TCL Q9M Pro 85寸（国补10499元，据雷科技）是目前万元出头价位参数很有竞争力的SQD-Mini LED产品，5000尼特+万象分区+300Hz超高刷，PS5和Xbox玩家会比较满意。如果想要98寸的超大屏体验，Q9M Pro 98寸国补价18499元（据雷科技），是目前少数把SQD-Mini LED做到98寸且价格在2万以内的产品。海信E7S Pro 100寸（20999元，京东商品页）也值得一看，玲珑真彩技术+330Hz高刷。')

    add_para(doc, '预算30000元以上，追求画质天花板', bold=True, size=12)
    add_para(doc, '海信UX 2026款是这个预算下的首选。10000尼特+43008分区+玲珑4芯+信芯H7 Pro，画质参数在国产电视里排名靠前。85寸34999元，100寸54999元（京东自营旗舰店），适合影音发烧友和专业创作人群。')

    add_para(doc, '特殊需求快速匹配', bold=True, size=12)
    for line in ['- 有孩家庭关注护眼：创维的莱茵TÜV双认证+ACD零频闪，三品牌里护眼投入力度很大',
                 '- 反感开机广告：TCL灵控系统3.0明确无开机广告，系统流畅度口碑不错',
                 '- 激光电视超大屏：海信激光电视全球份额53.5%（Omdia数据），产品成熟度高']:
        add_para(doc, line)

    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_场景决策图.png'),
        '按使用场景选品牌，一看就懂')

    # Conclusion
    add_heading(doc, '总结', level=2)

    add_para(doc, '海信、TCL、创维三家走的技术路数完全不同。海信押注画质技术深水区，RGB-Mini LED+自研芯片是硬核路线；TCL靠华星光电的面板产业链优势把SQD-Mini LED做到技术普惠，2026年一口气推出T7M Pro、Q9M Pro、T7M Ultra三款SQD新品，Mini LED全球出货暴增153.3%说明市场认了；创维选了一条差异化道路，用壁纸设计和护眼认证打动注重家居美学的消费群体。不存在"最好"的牌子，只有"更适合你"的选择。618加上国补的双重优惠快到了，按客厅尺寸、主要使用场景和预算上限来决定就好，别被品牌营销带着跑。')

    # Footer
    add_para(doc, '', size=6)
    add_para(doc, '本文配图均为 NotebookLM 生成信息图，数据来源见各图注。', size=9, color=RGBColor(0x99, 0x99, 0x99))
    add_para(doc, '发布时间建议：晚 8-10 点（头条流量较高时段）', size=9, color=RGBColor(0x99, 0x99, 0x99))

    out_path = os.path.join(OUTPUTS_DIR, '海信TCL创维电视横评_头条终稿_修订版.docx')
    doc.save(out_path)
    print(f'Word doc saved: {out_path}')
    return out_path


if __name__ == '__main__':
    build_doc()
