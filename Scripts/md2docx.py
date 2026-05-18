#!/usr/bin/env python3
"""
通用 Markdown → Word 文档转换器
读取 T6_final_头条.md，解析标题/段落/图片/粗体，输出 .docx
彻底避免在 Python 代码中硬编码中文内容导致的引号冲突问题

用法：
  python3 Scripts/md2docx.py                          # 默认读 outputs/T6_final_头条.md
  python3 Scripts/md2docx.py outputs/xxx.md           # 指定输入
  python3 Scripts/md2docx.py outputs/xxx.md out.docx  # 指定输出
"""

import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def main():
    if len(sys.argv) < 2:
        md_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                               'outputs', 'T6_final_头条.md')
    else:
        md_path = sys.argv[1]

    if len(sys.argv) >= 3:
        out_path = sys.argv[2]
    else:
        base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        # 从 md 文件名推断 docx 文件名
        basename = os.path.splitext(os.path.basename(md_path))[0]
        if 'T6' in basename or '头条' in basename:
            basename = basename.replace('T6_final_', '').replace('头条', '终稿')
        out_path = os.path.join(base, 'outputs', f'{basename}.docx')

    if not os.path.exists(md_path):
        print(f'错误: 找不到 {md_path}')
        sys.exit(1)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    outputs_dir = os.path.join(base_dir, 'outputs')

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Microsoft YaHei'
    style_normal.font.size = Pt(11)
    style_normal.paragraph_format.line_spacing = 1.8
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    def add_run(paragraph, text, bold=False, size=Pt(11), color=None):
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = 'Microsoft YaHei'
        run.font.size = size
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if color:
            run.font.color.rgb = color
        return run

    def add_text_with_bold(p, text, default_size=Pt(11)):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                add_run(p, part[2:-2], bold=True, size=default_size)
            else:
                add_run(p, part, bold=False, size=default_size)

    def add_image(filename, width_inches=5.8):
        path = os.path.join(outputs_dir, filename)
        if not os.path.exists(path):
            path = os.path.join(os.path.dirname(md_path), filename)
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run()
            run.add_picture(path, width=Inches(width_inches))
        else:
            p = doc.add_paragraph(f'[图片缺失: {filename}]')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 跳过 HTML 注释行
        if line.startswith('[//]: #'):
            i += 1
            continue

        # 空行跳过
        if not line.strip():
            i += 1
            continue

        # 图片 ![alt](filename)
        img_match = re.match(r'^!\[.*?\]\(([^)]+)\)', line)
        if img_match:
            add_image(img_match.group(1))
            i += 1
            continue

        # H1 标题
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(16)
            add_run(p, line[2:].strip(), bold=True, size=Pt(18),
                    color=RGBColor(0x1a, 0x1a, 0x1a))
            i += 1
            continue

        # H2 标题
        if line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            add_run(p, line[3:].strip(), bold=True, size=Pt(14),
                    color=RGBColor(0x1a, 0x1a, 0x1a))
            i += 1
            continue

        # H3 标题
        if line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            add_run(p, line[4:].strip(), bold=True, size=Pt(12))
            i += 1
            continue

        # 水平线
        if line.strip() == '---':
            i += 1
            continue

        # 普通段落（支持粗体混合）
        text = line.strip()
        p = doc.add_paragraph()
        add_text_with_bold(p, text)
        i += 1

    doc.save(out_path)
    size_mb = os.path.getsize(out_path) / 1024 / 1024
    print(f'Word文档已生成: {out_path}')
    print(f'文件大小: {size_mb:.1f} MB')


if __name__ == '__main__':
    main()
