#!/usr/bin/env python3
"""Generate Word doc from markdown with embedded images for TV brand review article."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUTS_DIR = os.path.join(BASE_DIR, "outputs")

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return h

def add_para(doc, text, bold=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p

def add_rich_para(doc, segments, size=11, space_after=6, align=None):
    """segments: list of (text, bold, color) tuples"""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold, color in segments:
        run = p.add_run(text)
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color
    return p

def add_image_with_caption(doc, img_path, caption, max_width=6.0):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(max_width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cap.add_run(caption)
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run2.italic = True

def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ===== 标题 =====
    title = doc.add_heading('海信、TCL、创维电视怎么选？三分钟看懂，别再交智商税', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

    doc.add_paragraph()  # spacer

    # ===== 封面图 =====
    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_封面图.png'),
        '2026年618海信、TCL、创维三大品牌电视选购横评')

    doc.add_paragraph()

    # ===== 开头 =====
    add_para(doc, '买电视这事儿，一句话就能说清楚：要画质选海信，要性价比和游戏体验选TCL，要颜值和护眼选创维。2025年中国电视零售量跌到2763万台（奥维睿沃数据），近十年新低，但Mini LED占比已经突破80%（洛图科技数据），"量减质升"的趋势摆在那儿。今天拿数据说话，把三大品牌各自的路数掰扯明白。', size=11)

    # ===== 市场格局 =====
    add_heading(doc, '电视市场格局：四家国产品牌吃掉80%以上份额', level=2)

    add_para(doc, '2025年全球彩电出货量大约2.64亿台（奥维睿沃+新浪财经数据），同比基本持平。中国市场这边就比较刺激了——零售量2763万台，出货量3289.5万台（洛图科技），两个数字口径不同，零售量是消费者实际掏钱买的，出货量是厂商发给渠道的，出货大于零售很正常。但格局已经定了：海信、TCL、小米、创维四家吞了超过80%的份额，前八大品牌占96%。长虹、海尔、康佳三家出货量加一块儿才103万台，同比还掉了12.2%，第二阵营在加速萎缩。三星、LG在中国市场更是持续边缘化，基本在慢慢退出家电业务（新浪财经报道）。', size=11)

    add_para(doc, '不过总量跌不代表没机会。大屏化趋势很明显，98寸和100寸的销量份额已经到2.8%左右，同比翻倍。换机周期拉长到10年以上（奥维睿沃数据），中国百户保有量超过120台，市场高度饱和。这种存量博弈下，Mini LED在2025年618期间占比突破80%（洛图科技数据），从"高端选项"变成了"标配"。我最近帮几个朋友挑电视，5000块就能买到85寸Mini LED，放两年前谁敢想。', size=11)

    # ===== 技术路线对比图 =====
    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_技术路线对比图.png'),
        '三大品牌核心技术路线对比（数据来源：品牌官方及京东商品页，2026年5月）')

    # ===== 海信 =====
    add_heading(doc, '海信：RGB-Mini LED画质天花板路线', level=2)

    add_para(doc, '先看海信的财务底盘。2025年年报显示，海信视像营收576.79亿元，同比微降1.45%，但归母净利润24.54亿元，反而涨了9.24%。营收降利润升，摆明了海信在主动调结构——宁可少卖也不打价格战，走"利润优先"。', size=11)

    add_para(doc, '技术层面，海信的杀手锏是RGB-Mini LED。旗舰UX 2026款搞了个"玲珑4芯真彩背光"，在红、绿之外加了颗青色（Cyan）发光芯片，填补了480-520nm波段的显示空白（量子位报道）。配合信芯AI H7 Pro画质芯片，这台机器堆到了10000尼特峰值亮度、43008个背光分区、110% BT.2020色域覆盖。京东自营旗舰店定价：85寸34999元，100寸54999元，冲着高端家庭和专业影音发烧友去的。黑曜屏Ultra反射率只有1.28%，强光下画面照样通透。帝瓦雷6.2.2声道音响系统在这个价位算领先配置，还支持2D转3D和Mac模式。', size=11)

    add_para(doc, '海信还有一条独门产品线——激光电视。Omdia数据显示（2022年统计），海信激光电视全球份额53.5%，8K激光电视属于首发量产级别。想要100寸以上的超大显示面积，又对环境光控制有条件的话，激光电视的护眼和尺寸优势值得认真考虑。', size=11)

    add_para(doc, 'RGB-Mini LED路线的优缺点：', bold=True, size=11)
    add_para(doc, '优点：', bold=True, size=11, color=RGBColor(0x22, 0x8B, 0x22))
    add_para(doc, '• 画质参数领先：10000尼特+43008分区的组合，对比度和HDR表现优势明显', size=11)
    add_para(doc, '• 自研画质芯片加持：信芯H7 Pro是国产品牌里少有的独立画质芯片，实时调校能力强', size=11)
    add_para(doc, '• 色域覆盖广：110% BT.2020配合青色芯片创新，色彩还原度高于传统量子点方案', size=11)
    add_para(doc, '不足：', bold=True, size=11, color=RGBColor(0xCC, 0x33, 0x33))
    add_para(doc, '• 价格门槛高：旗舰85寸34999元，对大多数家庭预算压力不小', size=11)
    add_para(doc, '• 激光电视受限：需要专业幕布和遮光环境，安装条件苛刻', size=11)
    add_para(doc, '• 入门款竞争力一般：海信中低端性价比不如TCL同价位产品', size=11)

    # ===== TCL =====
    add_heading(doc, 'TCL：QD-Mini LED性价比+游戏路线', level=2)

    add_para(doc, 'TCL的财务数据很炸。2025年年报显示，TCL电子营收1145.83亿港元，同比涨15.4%，归母净利润24.95亿港元，同比暴增41.8%。增长主要靠Mini LED产品线爆发——前三季度Mini LED出货量224万台，同比增153.3%（东方财富研报），全球占比11.4%，国内占比20.6%。北美市场Mini LED增速384.5%，欧洲124.1%。东方财富研报还提到，海信+TCL在600美元以上高端市场的合计份额，从2023年的22%飙到2025年一季度的39%。', size=11)

    add_para(doc, 'TCL的底层优势是"屏厂一体化"。旗下华星光电是全球面板大厂，从面板到整机一条龙自家产线，成本控制在国产品牌里很有底气。代表产品85S78A，85寸定价5999元（2026年85寸电视性价比排行榜），高端VA屏，原生对比度7000:1，4K 144Hz刷新率支持全通道288Hz，关键是明确标注无开机广告。这个价位给到这个配置，85寸市场里确实能打。', size=11)

    add_para(doc, '98鹤7Pro 26款定价18234元（京东商品页，含国补减1500元），万象分区+蝶翼华曜屏+360Hz刷新率。鹏7 Plus 85寸定价12499元（京东商品页），360Hz高刷+94% DCI-P3色域。旗舰X11L用的是SQD-Mini LED，ZeroBorder无边框设计，机身最薄处不到0.8英寸，B&O定制音响。', size=11)

    add_para(doc, 'QD-Mini LED路线的优缺点：', bold=True, size=11)
    add_para(doc, '优点：', bold=True, size=11, color=RGBColor(0x22, 0x8B, 0x22))
    add_para(doc, '• 性价比突出：5999元买85寸Mini LED，华星光电面板+无开机广告，价格有吸引力', size=11)
    add_para(doc, '• 游戏体验好：288Hz/360Hz高刷+FreeSync认证，输入延迟低，PS5和Xbox适配成熟', size=11)
    add_para(doc, '• 全球化布局强：北美、欧洲增速亮眼，产品迭代快，供应链稳定', size=11)
    add_para(doc, '不足：', bold=True, size=11, color=RGBColor(0xCC, 0x33, 0x33))
    add_para(doc, '• 画质芯片相对弱：没有独立自研画质芯片，图像调校更多依赖面板素质和通用算法', size=11)
    add_para(doc, '• 高端线差异化不够：旗舰X11L在画质天花板层面跟海信UX还有差距', size=11)
    add_para(doc, '• 品牌溢价低：高端市场的品牌认知度不如海信', size=11)

    # ===== 创维 =====
    add_heading(doc, '创维：壁纸电视美学+护眼路线', level=2)

    add_para(doc, '创维的财务状况有点意思。2025年年报显示，创维集团营收703.24亿元，同比涨了8.2%，但归母净利润只有3.56亿元，反而掉了37.3%。利润下滑主要跟新能源等第二曲线的投入有关，电视主业整体处于战略转型期。', size=11)

    add_para(doc, '创维的核心策略很聪明——避开跟海信、TCL的参数军备竞赛，走"家居美学+护眼"的差异化路线。代表产品A7H Pro 85寸，京东商品页标价9763.52元（2026年5月京东自营价，含国补后约8299元）。这台机器的核心卖点不在参数堆砌，而在设计感——约3-4cm超薄壁纸设计，贴墙安装后跟挂画差不多。画境屏用AG/LR低反技术，99% DCI-P3色域，ΔE≤0.6的专业级色准，1344个Mini LED分区，3800尼特峰值亮度（2026年85寸电视性价比排行榜，新浪财经）。哈曼卡顿双Soundbar音响加分不少，莱茵TÜV低蓝光和无频闪双认证对于有孩家庭来说实打实的健康考量。4K 150Hz刷新率支持倍频300Hz，日常追剧和看球够用。', size=11)

    add_para(doc, '旗舰款A10H 100寸定价25599元（京东商品页），6120分区SQD Mini LED，跟京东方合作的面板方案，极黑广角类纸屏。酷开AIOS系统接入了DeepSeek-R1大模型（京东商品页），智能交互上有不错的创新。', size=11)

    add_para(doc, '壁纸Mini LED路线的优缺点：', bold=True, size=11)
    add_para(doc, '优点：', bold=True, size=11, color=RGBColor(0x22, 0x8B, 0x22))
    add_para(doc, '• 家居美学突出：3-4cm超薄壁纸设计是国产品牌里少有的"去家电化"产品，适合讲究客厅整体风格的家庭', size=11)
    add_para(doc, '• 护眼配置到位：莱茵TÜV双认证+ACD零频闪，有孩家庭的健康考量比较充分', size=11)
    add_para(doc, '• 色准专业级：ΔE≤0.6在万元价位段属于领先水平，对色彩敏感的用户会比较满意', size=11)
    add_para(doc, '不足：', bold=True, size=11, color=RGBColor(0xCC, 0x33, 0x33))
    add_para(doc, '• 性价比不占优：9000元上下的价位在85寸市场面临TCL和海信的强力竞争', size=11)
    add_para(doc, '• 画质天花板不够高：3800尼特和1344分区在中高端算主流水平，跟海信UX和TCL旗舰有明显差距', size=11)
    add_para(doc, '• 品牌盈利承压：净利润掉了37.3%，研发投入可能受限，长期产品力存在不确定性', size=11)

    # ===== 参数横评 =====
    add_heading(doc, '85寸主力机型参数横评', level=2)

    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_参数对比图.png'),
        '85寸主力机型核心参数对比（数据来源：京东商品页及2026年85寸电视性价比排行榜）')

    add_para(doc, '四款85寸主力机型放一起看看：', size=11)

    table_headers = ['参数', '创维 A7H Pro', 'TCL 85S78A', '雷鸟 85R69A', '海信 85E3Q']
    table_rows = [
        ['参考价', '¥9,763（含国补约¥8,299）', '¥5,999', '¥5,799', '¥4,999'],
        ['背光/分区', 'Mini LED 1344分区', '强控光', 'Mini LED 400+分区', 'AI色彩增强'],
        ['峰值亮度', '3800尼特', '未标注', '1300尼特', '未标注'],
        ['刷新率', '4K 150Hz（倍频300Hz）', '4K 144Hz（全通道288Hz）', '4K 144Hz（支持288Hz）', '144Hz+MEMC'],
        ['色域', '99% DCI-P3', '未标注', '未标注', '未标注'],
        ['色准', 'ΔE≤0.6', '未标注', '未标注', '未标注'],
        ['音响', '哈曼卡顿双Soundbar', '安桥2.1声道50W', '安桥2.1+独立低音炮', '未标注'],
        ['特色功能', '超薄壁纸+莱茵护眼', '无开机广告', 'FreeSync Premium Pro', 'AI语音+低蓝光护眼'],
    ]
    add_table(doc, table_headers, table_rows)

    source_note = doc.add_paragraph()
    source_note.paragraph_format.space_before = Pt(4)
    run = source_note.add_run('数据来源：价格据京东商品页（2026年5月采集）和2026年85寸电视性价比排行榜（新浪财经，2026年2月）；产品参数据品牌官方和京东商品页。')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.add_paragraph()
    add_para(doc, '几个关键差异一目了然：TCL 85S78A在5999元给了288Hz高刷+无开机广告，性价比非常有竞争力；创维A7H Pro的1344分区、3800尼特和色准ΔE≤0.6明显领先同价位；雷鸟85R69A拿Mini LED 400+分区打到5799元，属于5K价位的强力竞争者；海信85E3Q定价4999元，85寸入门市场的守门员。', size=11)

    # ===== 场景选购 =====
    add_heading(doc, '按场景选购：搞清需求比纠结品牌重要', level=2)

    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_价格段选购图.png'),
        '2026年618电视按预算对号入座（价格来源：京东，2026年5月）')

    # 6000以内
    add_para(doc, '预算6000元以内，想买85寸大屏', bold=True, size=11)
    add_para(doc, '海信85E3Q（4999元，京东商品页）和雷鸟85R69A（5799元，京东商品页）是这个价位的主力。海信的优势在画质调校积累和售后体系，雷鸟的优势在Mini LED 400+分区在这个价位很有竞争力。更在乎画质一致性选海信，追求参数性价比选雷鸟。', size=11)

    # 6000-10000
    add_para(doc, '预算6000到10000元，兼顾画质和日常', bold=True, size=11)
    add_para(doc, 'TCL 85S78A（5999元）在这个价位段确实能打，无开机广告+288Hz高刷+安桥音响，日常追剧和偶尔打游戏都能覆盖。预算再往上走，创维A7H Pro（含国补约8299元）画质和美学更进一步，1344分区+哈曼卡顿的组合适合对音画品质有较高要求的家庭。', size=11)

    # 10000-20000
    add_para(doc, '预算10000到20000元，追求大屏游戏', bold=True, size=11)
    add_para(doc, 'TCL 98鹤7Pro（18234元，京东商品页）是目前少数把98寸做到20000元以内的Mini LED产品，360Hz刷新率+FreeSync，PS5和Xbox玩家会比较满意。海信E7S Pro 100寸（20999元，京东商品页）也值得一看，玲珑真彩技术+330Hz高刷。', size=11)

    # 30000+
    add_para(doc, '预算30000元以上，追求画质天花板', bold=True, size=11)
    add_para(doc, '海信UX 2026款是这个预算下的首选。10000尼特+43008分区+玲珑4芯+信芯H7 Pro，画质参数在国产电视里排名靠前。85寸34999元，100寸54999元（京东自营旗舰店），适合影音发烧友和专业创作人群。', size=11)

    # 特殊需求
    add_para(doc, '特殊需求快速匹配', bold=True, size=11)
    add_para(doc, '• 有孩家庭关注护眼：创维的莱茵TÜV双认证+ACD零频闪，三品牌里护眼投入力度很大', size=11)
    add_para(doc, '• 反感开机广告：TCL明确无开机广告，系统流畅度口碑不错', size=11)
    add_para(doc, '• 激光电视超大屏：海信激光电视全球份额53.5%（Omdia数据），产品成熟度高', size=11)

    # 场景决策图
    add_image_with_caption(doc,
        os.path.join(OUTPUTS_DIR, 'T5_场景决策图.png'),
        '按使用场景选品牌，一看就懂')

    # ===== 总结 =====
    add_heading(doc, '总结', level=2)

    add_para(doc, '海信、TCL、创维三家走的技术路数完全不同。海信押注画质技术深水区，RGB-Mini LED+自研芯片是硬核路线；TCL靠华星光电的面板产业链优势把性价比做到极致，Mini LED全球出货暴增153.3%说明市场认了；创维选了一条差异化道路，用壁纸设计和护眼认证打动注重家居美学的消费群体。不存在"最好"的牌子，只有"更适合你"的选择。618加上国补的双重优惠快到了，按客厅尺寸、主要使用场景和预算上限来决定就好，别被品牌营销带着跑。', size=11)

    # ===== 页脚 =====
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('本文配图均为 NotebookLM 生成信息图，数据来源见各图注')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.italic = True

    # Save
    out_path = os.path.join(BASE_DIR, 'outputs', '海信TCL创维电视横评_头条终稿.docx')
    doc.save(out_path)
    print(f'Word document saved to: {out_path}')

if __name__ == '__main__':
    main()
